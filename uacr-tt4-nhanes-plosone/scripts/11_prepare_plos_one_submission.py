from __future__ import annotations

import csv
import math
import re
from datetime import datetime
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_TITLE = (
    "Urinary albumin-to-creatinine ratio, total thyroxine, and mortality risk "
    "among U.S. adults: evidence from NHANES 2007\u20132012"
)
SHORT_TITLE = "UACR, total thyroxine, and mortality"

AUTHOR_NAMES = "Haoyang Li; Xiang Li; Xuefeng Shi"
AUTHOR_LINE = "Haoyang Li1\u2020, Xiang Li2\u2020, Xuefeng Shi3*"
EQUAL_CONTRIBUTION = (
    "\u2020Haoyang Li and Xiang Li contributed equally to this work and share first authorship."
)
AFFILIATIONS = [
    "1Department of Cardiovascular Medicine, Qinghai University Affiliated Hospital, "
    "No. 29 Tongren Road, Chengxi District, Xining, Qinghai 810000, China.",
    "2Department of Thyroid and Breast Surgery, Qinghai University Affiliated Hospital, "
    "No. 29 Tongren Road, Chengxi District, Xining, Qinghai 810000, China.",
    "3Department of Respiratory and Critical Care Medicine, Qinghai Provincial People\u2019s Hospital, "
    "No. 2 Gonghe Road, Chengdong District, Xining, Qinghai 810000, China.",
]
CORRESPONDING_AUTHOR = "Xuefeng Shi"
CORRESPONDING_DEPARTMENT = "Department of Respiratory and Critical Care Medicine"
CORRESPONDING_INSTITUTION = "Qinghai Provincial People\u2019s Hospital"
CORRESPONDING_ADDRESS = "No. 2 Gonghe Road, Chengdong District, Xining, Qinghai 810000, China"
CORRESPONDING_EMAIL = "shixuefeng128@163.com"
AUTHOR_DETAILS = [
    (
        "Haoyang Li",
        "Department of Cardiovascular Medicine",
        "Qinghai University Affiliated Hospital",
        "No. 29 Tongren Road, Chengxi District, Xining, Qinghai 810000, China",
        "Lihaoyang369@outlook.com",
        "0009-0006-0399-9464",
    ),
    (
        "Xiang Li",
        "Department of Thyroid and Breast Surgery",
        "Qinghai University Affiliated Hospital",
        "No. 29 Tongren Road, Chengxi District, Xining, Qinghai 810000, China",
        "lixiang081601@outlook.com",
        "0009-0008-7853-5415",
    ),
    (
        "Xuefeng Shi",
        CORRESPONDING_DEPARTMENT,
        CORRESPONDING_INSTITUTION,
        CORRESPONDING_ADDRESS,
        CORRESPONDING_EMAIL,
        "0000-0002-4694-8759",
    ),
]
AUTHOR_CONTRIBUTIONS = (
    "H.L. and X.L. contributed equally to this work and share first authorship. "
    "H.L. conceived the study, designed the analytical workflow, performed data processing and "
    "statistical analyses, interpreted the results, prepared tables and figures, and drafted the "
    "manuscript. X.L. contributed to study design, clinical interpretation, data review, and critical "
    "manuscript revision. X.S. supervised the study, provided clinical oversight, reviewed and revised "
    "the manuscript critically, and is the corresponding author. All authors read and approved the "
    "final manuscript. [AUTHORS TO CONFIRM FOR THIS STUDY.]"
)
CREDIT_STATEMENT = (
    "CRediT author statement: Conceptualization: H.L.; Data curation: H.L.; Formal analysis: H.L.; "
    "Methodology: H.L., X.L.; Software: H.L.; Validation: H.L., X.L.; Visualization: H.L.; "
    "Supervision: X.S.; Writing - original draft: H.L.; Writing - review and editing: X.L., X.S. "
    "[AUTHORS TO CONFIRM FOR THIS STUDY.]"
)
FUNDING_STATEMENT = (
    "This study was supported by the National Excellent Young Physician Program "
    "(Document No. 2024[41]). The funder had no role in study design, data collection and analysis, "
    "decision to publish, or preparation of the manuscript."
)
COMPETING_INTERESTS_STATEMENT = "The authors have declared that no competing interests exist."

PLOS_GUIDELINES_URL = "https://journals.plos.org/plosone/s/submission-guidelines"
PLOS_DATA_URL = "https://journals.plos.org/plosone/s/data-availability"
STROBE_URL = (
    "https://strobe-statement.org/fileadmin/Strobe/uploads/checklists/"
    "STROBE_checklist_v4_cohort.pdf"
)
NHANES_PUBLIC_URL = "https://wwwn.cdc.gov/nchs/nhanes/continuousnhanes/default.aspx"
NHANES3_URL = "https://wwwn.cdc.gov/nchs/nhanes/nhanes3/datafiles.aspx"
NHANES_ETHICS_URL = "https://www.cdc.gov/nchs/nhanes/about/erb.html"
MORTALITY_URL = "https://www.cdc.gov/nchs/data-linkage/mortality-public.htm"
OPENGWAS_URL = "https://gwas.mrcieu.ac.uk/"
ETHICS_STATEMENT = (
    "This study was a secondary analysis of publicly available, de-identified data from NHANES and "
    "NHANES III. NHANES protocols were reviewed and approved by the NCHS Research Ethics Review Board, "
    "and written informed consent was obtained from participants. Information on NHANES ethics review "
    f"is available from NCHS: {NHANES_ETHICS_URL}. The present study involved no direct participant "
    "contact and used only public-use files; therefore, additional institutional review board approval "
    "was not required."
)


def find_project_root(start: Path | None = None) -> Path:
    current = (start or Path.cwd()).resolve()
    for candidate in [current, *current.parents]:
        if (candidate / "config" / "analysis_plan.yaml").exists():
            return candidate
    raise RuntimeError("Could not find project root containing config/analysis_plan.yaml.")


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.rstrip() + "\n", encoding="utf-8")


def write_csv(path: Path, rows: list[dict[str, str]], fields: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields)
        writer.writeheader()
        writer.writerows(rows)


def fmt_num(value: float, digits: int = 3) -> str:
    if pd.isna(value):
        return ""
    return f"{value:.{digits}f}"


def fmt_p(value: float) -> str:
    if pd.isna(value):
        return ""
    if value < 0.001:
        return "<0.001"
    return f"{value:.3f}"


def fmt_estimate(row: pd.Series, kind: str = "beta") -> str:
    estimate = row["beta"] if kind == "beta" else row["hr"]
    return f"{estimate:.3f} ({row['ci_low']:.3f}-{row['ci_high']:.3f})"


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 8) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_line_numbers(section) -> None:
    sect_pr = section._sectPr
    line_number = sect_pr.find(qn("w:lnNumType"))
    if line_number is None:
        line_number = OxmlElement("w:lnNumType")
        sect_pr.append(line_number)
    line_number.set(qn("w:countBy"), "1")
    line_number.set(qn("w:restart"), "continuous")


def set_cell_margins(cell, top: int = 40, start: int = 50, bottom: int = 40, end: int = 50) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure_document(document: Document, *, line_numbers: bool = True) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_after = Pt(0)

    for style_name, size in (("Title", 14), ("Heading 1", 12), ("Heading 2", 12), ("Heading 3", 12)):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(0)

    if "Caption" not in styles:
        caption = styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Caption"]
    caption.font.name = "Times New Roman"
    caption.font.size = Pt(10)
    caption.font.italic = False
    caption.paragraph_format.line_spacing = 1
    caption.paragraph_format.space_after = Pt(3)

    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        if line_numbers:
            add_line_numbers(section)
        add_page_number(section.footer.paragraphs[0])


def add_paragraph(document: Document, text: str = "", *, bold_lead: str | None = None):
    paragraph = document.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        paragraph.add_run(bold_lead).bold = True
        paragraph.add_run(text[len(bold_lead) :])
    else:
        paragraph.add_run(text)
    return paragraph


def add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_table(document: Document, headers: list[str], rows: list[list[str]], *, font_size: int = 8):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, size=font_size)
        shade_cell(table.rows[0].cells[index], "D9EAF7")
        set_cell_margins(table.rows[0].cells[index])
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            set_cell_text(row.cells[index], value, size=font_size)
            set_cell_margins(row.cells[index])
    document.add_paragraph()
    return table


def first_row(data: pd.DataFrame, **criteria) -> pd.Series:
    selected = data.copy()
    for column, value in criteria.items():
        selected = selected[selected[column] == value]
    if len(selected) != 1:
        raise RuntimeError(f"Expected one row for {criteria}, found {len(selected)}.")
    return selected.iloc[0]


def build_main_table_2(table2: pd.DataFrame) -> list[list[str]]:
    rows: list[list[str]] = []
    for label, exposure_type, contrast in (
        ("Natural-log UACR, per 1-unit increase", "log_UACR", "LOG_UACR"),
        ("UACR 30-300 vs <30 mg/g", "UACR clinical category", "UACR_CLINICAL_CATEGORY30-300"),
        ("UACR >=300 vs <30 mg/g", "UACR clinical category", "UACR_CLINICAL_CATEGORY>=300"),
    ):
        values = [label]
        for model in ("Model 1", "Model 2", "Model 3"):
            row = first_row(
                table2,
                outcome="TT4",
                exposure_type=exposure_type,
                contrast=contrast,
                model=model,
            )
            values.append(f"{fmt_estimate(row)}; P={fmt_p(row['p_value'])}")
        rows.append(values)

    trends = []
    for model in ("Model 1", "Model 2", "Model 3"):
        row = first_row(
            table2,
            outcome="TT4",
            exposure_type="UACR clinical category",
            contrast="UACR_CLINICAL_CATEGORY30-300",
            model=model,
        )
        trends.append(f"P for trend={fmt_p(row['p_trend'])}")
    rows.append(["Clinical-category ordinal trend", *trends])
    return rows


def build_main_table_3(mortality: pd.DataFrame) -> list[list[str]]:
    rows: list[list[str]] = []
    labels = {
        "all_cause_mortality": "All-cause mortality",
        "cardiovascular_mortality": "Cardiovascular mortality",
    }
    for outcome in ("all_cause_mortality", "cardiovascular_mortality"):
        for exposure in ("LOG_UACR", "TT4"):
            h1 = first_row(mortality, outcome=outcome, exposure=exposure, contrast=exposure, model="H1")
            full = first_row(mortality, outcome=outcome, exposure=exposure, contrast=exposure, model="Full")
            exposure_label = "Natural-log UACR" if exposure == "LOG_UACR" else "TT4, per 1-unit increase"
            rows.append(
                [
                    labels[outcome],
                    exposure_label,
                    fmt_estimate(h1, "hr"),
                    fmt_p(h1["p_value"]),
                    fmt_estimate(full, "hr"),
                    fmt_p(full["p_value"]),
                    str(int(full["events"])),
                ]
            )
    return rows


def statement_texts() -> dict[str, str]:
    data_availability = f"""# Data Availability

The study used publicly available, de-identified data. NHANES 2007-2012 data and documentation are available from the U.S. Centers for Disease Control and Prevention (CDC), National Center for Health Statistics (NCHS): {NHANES_PUBLIC_URL}. NHANES III data and documentation are available from NCHS: {NHANES3_URL}. Public-use linked mortality files and documentation are available from NCHS: {MORTALITY_URL}. Summary-level GWAS data used for the Supplementary exploratory genetic analysis were retrieved from OpenGWAS: {OPENGWAS_URL}.

The analysis scripts, configuration files, derived analytic tables, figure source data, and logs supporting the reported results will be deposited at [INSERT PUBLIC REPOSITORY DOI OR PERSISTENT URL BEFORE SUBMISSION]. The repository record should exclude any credentials or API tokens.

## Author action required

- Insert a durable public repository DOI or persistent URL before submission.
- Confirm that the repository contains the version of the code and derived outputs used for the submitted manuscript.
- Do not upload `OPENGWAS_JWT` or any local credential files.
"""

    ethics = f"""# Ethics Statement

{ETHICS_STATEMENT}
"""

    competing = f"""# Competing Interests Statement

{COMPETING_INTERESTS_STATEMENT}
"""

    funding = f"""# Funding Statement

{FUNDING_STATEMENT}
"""
    return {
        "PLOS_ONE_data_availability_statement.md": data_availability,
        "PLOS_ONE_ethics_statement.md": ethics,
        "PLOS_ONE_competing_interests_statement.md": competing,
        "PLOS_ONE_funding_statement.md": funding,
    }


def add_corresponding_author_block(document: Document) -> None:
    add_paragraph(document, CORRESPONDING_AUTHOR)
    add_paragraph(document, CORRESPONDING_DEPARTMENT)
    add_paragraph(document, CORRESPONDING_INSTITUTION)
    add_paragraph(document, CORRESPONDING_ADDRESS)
    add_paragraph(document, f"Email: {CORRESPONDING_EMAIL}")


def make_title_page(root: Path) -> None:
    document = Document()
    configure_document(document, line_numbers=False)
    document.core_properties.title = f"Title page: {PROJECT_TITLE}"
    document.core_properties.subject = "PLOS ONE title page"
    document.core_properties.author = AUTHOR_NAMES

    document.add_heading("Title Page", level=1)
    document.add_heading("Title", level=2)
    add_paragraph(document, PROJECT_TITLE)
    document.add_heading("Running title", level=2)
    add_paragraph(document, SHORT_TITLE)
    document.add_heading("Authors", level=2)
    add_paragraph(document, AUTHOR_LINE)
    add_paragraph(document, EQUAL_CONTRIBUTION)
    add_paragraph(document, "*Corresponding author.")
    document.add_heading("Affiliations", level=2)
    for affiliation in AFFILIATIONS:
        add_paragraph(document, affiliation)
    document.add_heading("Corresponding author", level=2)
    add_corresponding_author_block(document)
    document.add_heading("Author information", level=2)
    for name, department, institution, address, email, orcid in AUTHOR_DETAILS:
        add_paragraph(
            document,
            f"{name}\n{department}\n{institution}\n{address}\nEmail: {email}\nORCID: {orcid}",
        )
    document.add_heading("Authors' contributions", level=2)
    add_paragraph(document, AUTHOR_CONTRIBUTIONS)
    add_paragraph(document, CREDIT_STATEMENT)
    document.add_heading("Funding", level=2)
    add_paragraph(document, FUNDING_STATEMENT)
    document.add_heading("Competing interests", level=2)
    add_paragraph(document, COMPETING_INTERESTS_STATEMENT)
    document.add_heading("Ethics approval and consent to participate", level=2)
    add_paragraph(document, ETHICS_STATEMENT)

    output = root / "manuscript" / "PLOS_ONE_title_page.docx"
    document.save(output)


def make_main_manuscript(
    root: Path,
    table1: pd.DataFrame,
    table2: pd.DataFrame,
    mortality: pd.DataFrame,
    followup_summary: str,
) -> None:
    document = Document()
    configure_document(document, line_numbers=True)
    document.core_properties.title = PROJECT_TITLE
    document.core_properties.subject = "PLOS ONE submission draft"
    document.core_properties.author = AUTHOR_NAMES

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(PROJECT_TITLE)
    add_paragraph(document, AUTHOR_LINE)
    add_paragraph(document, EQUAL_CONTRIBUTION)
    for affiliation in AFFILIATIONS:
        add_paragraph(document, affiliation)
    add_paragraph(
        document,
        f"Corresponding author: {CORRESPONDING_AUTHOR}; {CORRESPONDING_DEPARTMENT}; "
        f"{CORRESPONDING_INSTITUTION}; {CORRESPONDING_ADDRESS}; Email: {CORRESPONDING_EMAIL}",
    )
    add_paragraph(document, f"Short title: {SHORT_TITLE}")
    add_paragraph(document, "Article type: Research Article")
    add_paragraph(document, "Word count: [UPDATE AFTER FINAL EDITING]")

    document.add_heading("Abstract", level=1)
    abstract = (
        "Urinary albumin-to-creatinine ratio (UACR) is an established marker of kidney and vascular "
        "risk, but its association with thyroid hormone homeostasis is incompletely understood. We "
        "evaluated the association of UACR with total thyroxine (TT4) in 6487 adults from NHANES "
        "2007-2012 using survey-weighted models and linked 6484 eligible participants to public-use "
        "mortality data. Higher natural-log UACR was associated with higher TT4 after multivariable "
        "adjustment (beta=0.077, 95% CI 0.035-0.119; P=0.001). The association remained positive "
        "across prespecified sensitivity analyses and showed no evidence of non-linearity. During "
        "mortality follow-up, 897 all-cause deaths and 249 cardiovascular deaths were identified. "
        "Natural-log UACR was associated with all-cause mortality (HR=1.387, 95% CI 1.285-1.498; "
        "P<0.001) and cardiovascular mortality (HR=1.331, 95% CI 1.161-1.526; P<0.001). TT4 showed "
        "a modest association with all-cause mortality (HR=1.065, 95% CI 1.009-1.125; P=0.022) but "
        "not cardiovascular mortality (HR=1.108, 95% CI 0.978-1.255; P=0.108). In a harmonized "
        "NHANES III analysis, the UACR-TT4 association was not statistically replicated. These "
        "observational findings identify a stable association between UACR and TT4 in NHANES "
        "2007-2012 and reinforce UACR as a mortality risk marker. The TT4 finding requires "
        "confirmation in additional populations."
    )
    add_paragraph(document, abstract)
    add_paragraph(document, "Keywords: albuminuria; total thyroxine; mortality; NHANES; survey analysis")

    document.add_heading("Introduction", level=1)
    add_paragraph(
        document,
        "Urinary albumin-to-creatinine ratio (UACR) is widely used to quantify albuminuria and "
        "captures clinically relevant kidney and vascular risk. Even modest increases in UACR may "
        "reflect systemic vascular injury and are associated with adverse outcomes. [INSERT "
        "BACKGROUND REFERENCES]",
    )
    add_paragraph(
        document,
        "Kidney function and thyroid hormone homeostasis are physiologically connected through "
        "filtration, metabolism, protein handling, and iodine-related pathways. However, population-"
        "based evidence linking UACR to circulating thyroid indicators remains limited, and observed "
        "associations may vary across survey eras and assay platforms. [INSERT BACKGROUND REFERENCES]",
    )
    add_paragraph(
        document,
        "We therefore evaluated whether UACR was associated with total thyroxine (TT4) among adults "
        "in NHANES 2007-2012. We assessed the stability of this association across exposure "
        "definitions, restricted cubic spline models, and prespecified sensitivity analyses. We "
        "also examined all-cause and cardiovascular mortality using public-use linked mortality "
        "files. A harmonized NHANES III assessment and an exploratory genetic analysis were retained "
        "as Supplementary analyses to define the limits of interpretation.",
    )

    document.add_heading("Materials and methods", level=1)
    document.add_heading("Study design and data sources", level=2)
    add_paragraph(
        document,
        "We conducted an observational study using adults from the 2007-2008, 2009-2010, and "
        "2011-2012 cycles of the National Health and Nutrition Examination Survey (NHANES). NHANES "
        "uses a complex, multistage probability sampling design to represent the non-institutionalized "
        "U.S. population. The primary analysis was cross-sectional. A prospective mortality extension "
        "linked the discovery cohort to the public-use 2019 NCHS Linked Mortality Files, with follow-up "
        "through December 31, 2019. NHANES III was analyzed separately as a Supplementary replication "
        "assessment. An exploratory genetic analysis using OpenGWAS summary statistics was also "
        "restricted to the Supplementary Materials.",
    )
    add_paragraph(
        document,
        ETHICS_STATEMENT,
    )

    document.add_heading("NHANES 2007-2012 discovery cohort", level=2)
    add_paragraph(
        document,
        "Component files were merged by participant identifier (SEQN) within each survey cycle and "
        "then pooled. Participants were eligible if they were aged at least 18 years, were not "
        "pregnant, had UACR greater than zero, had non-missing common thyroid indicators used during "
        "cohort construction (thyroid-stimulating hormone [TSH], TT4, thyroglobulin antibody [TGAb], "
        "and thyroid peroxidase antibody [TPOAb]), were not identified as thyroid-medication users, "
        "and had complete prespecified cohort-construction covariates. Complete physical-activity "
        "information was required during cohort construction. The final discovery cohort included "
        "6487 participants. The complete-case approach and exclusions are documented in Supplementary "
        "Table S1.",
    )
    add_paragraph(
        document,
        "Cycle-specific thyroid-related weights were used. The 2007-2008 cycle used the 2-year mobile "
        "examination center weight (WTMEC2YR), whereas the 2009-2010 and 2011-2012 cycles used the "
        "thyroid subsample weight (WTSA2YR). The selected 2-year weight was divided by three to obtain "
        "the pooled 6-year analysis weight. Survey designs incorporated the pooled weight, primary "
        "sampling units (SDMVPSU), and strata (SDMVSTRA).",
    )

    document.add_heading("Exposure, thyroid outcome, and covariates", level=2)
    add_paragraph(
        document,
        "UACR was expressed in mg/g. The primary exposure was the natural logarithm of UACR. "
        "Secondary definitions were UACR quartiles and clinical categories of <30, 30-300, and "
        ">=300 mg/g. TT4 was the primary thyroid outcome. TSH, free triiodothyronine, free thyroxine, "
        "total triiodothyronine, and thyroglobulin were secondary outcomes; TGAb and TPOAb were "
        "exploratory thyroid-autoimmunity outcomes and were not used to define the main conclusion.",
    )
    add_paragraph(
        document,
        "Prespecified covariates included age, sex, race/ethnicity, education, poverty-income ratio "
        "(PIR), body mass index (BMI), smoking, alcohol use, diabetes, hypertension, estimated "
        "glomerular filtration rate (eGFR), and urinary iodine concentration (UIC). eGFR was "
        "calculated using the 2021 CKD-EPI creatinine equation. [INSERT CKD-EPI REFERENCE]",
    )

    document.add_heading("Survey-weighted thyroid analyses", level=2)
    add_paragraph(
        document,
        "Associations of UACR with TT4 were estimated using survey-weighted generalized linear models. "
        "Model 1 was unadjusted. Model 2 adjusted for age, sex, and race/ethnicity. Model 3 additionally "
        "adjusted for education, PIR, BMI, smoking, alcohol use, diabetes, hypertension, eGFR, and UIC. "
        "Continuous models estimated the TT4 difference per one-unit increase in natural-log UACR. "
        "Categorical models evaluated quartiles and clinical categories, with ordinal scores used for "
        "trend tests. False-discovery-rate correction was applied across thyroid outcomes.",
    )
    add_paragraph(
        document,
        "Restricted cubic spline models used knots at the 5th, 35th, 65th, and 95th percentiles of "
        "natural-log UACR. Sensitivity analyses excluded participants with eGFR <60 mL/min/1.73 m2, "
        "diabetes, hypertension, UACR >=300 mg/g, or UACR values outside the 1st-99th percentile range. "
        "A separate analysis was restricted to euthyroid participants, defined as TSH 0.45-4.50 and "
        "free thyroxine 0.60-1.60.",
    )

    document.add_heading("Mortality follow-up and Cox models", level=2)
    add_paragraph(
        document,
        "The discovery cohort was linked by SEQN to the public-use 2019 NCHS Linked Mortality Files. "
        "Participants were eligible for mortality analyses when ELIGSTAT equaled 1 and follow-up "
        "months from the mobile examination center examination (PERMTH_EXM) were available. Follow-up "
        "time was calculated as PERMTH_EXM divided by 12. All-cause mortality was defined using "
        "MORTSTAT. Cardiovascular mortality was defined among decedents using the public-use leading "
        "underlying-cause categories for heart disease or cerebrovascular disease (UCOD_LEADING values "
        "1 or 5).",
    )
    add_paragraph(
        document,
        "Survey-weighted Cox proportional-hazards models incorporated the pooled 6-year analysis "
        "weight, SDMVPSU, and SDMVSTRA. Mortality Model 1 adjusted for age, sex, and race/ethnicity. "
        "Model 2 additionally adjusted for education, PIR, BMI, smoking, and alcohol use. Model 3 "
        "additionally adjusted for diabetes, hypertension, eGFR, and UIC. Natural-log UACR was the "
        "primary mortality exposure. TT4 was evaluated as a secondary prognostic marker, including a "
        "per-standard-deviation sensitivity analysis. Prespecified sensitivity analyses excluded "
        "deaths occurring within the first two years of follow-up and restricted analyses to "
        "euthyroid participants. Proportional-hazards assumptions were evaluated diagnostically using "
        "non-weighted Cox models with the same Model 3 covariates because a validated direct cox.zph "
        "workflow is not available for survey-weighted Cox models.",
    )
    add_paragraph(
        document,
        "Secondary descriptive analyses evaluated four groups defined by UACR <30 or >=30 mg/g and "
        "TT4 below or within the weighted highest quartile. Multiplicative effect-modification terms "
        "were evaluated as secondary analyses. These analyses were not used to define the primary "
        "mortality conclusion.",
    )

    document.add_heading("Supplementary replication and exploratory genetic analyses", level=2)
    add_paragraph(
        document,
        "NHANES III was analyzed separately using harmonized exposure definitions and covariate blocks. "
        "NHANES III survey designs incorporated WTPFEX6, SDPPSU6, and SDPSTRA6. UIC was omitted from "
        "harmonized comparisons because it was unavailable for the NHANES III comparison. Details are "
        "reported in Supplementary Tables S4 and S5.",
    )
    add_paragraph(
        document,
        "An exploratory bidirectional genetic analysis used OpenGWAS summary statistics. Genetic "
        "instrument selection, available trait pairs, and instrument limitations are reported in "
        "Supplementary Tables S10-S12. This analysis was not used as evidence of a direct UACR-TT4 "
        "pathway.",
    )

    document.add_heading("Statistical software", level=2)
    add_paragraph(
        document,
        "Data construction used Python 3.12.7. Statistical analyses used R 4.5.3, including survey "
        "version 4.5 and survival version 3.8.6. OpenGWAS access used ieugwasr version 1.1.0. "
        "Two-sided P values <0.05 were considered nominally significant. Analysis scripts, derived "
        "tables, figure source data, and logs will be deposited at [INSERT PUBLIC REPOSITORY DOI OR "
        "PERSISTENT URL BEFORE SUBMISSION].",
    )

    document.add_heading("Results", level=1)
    document.add_heading("Study population", level=2)
    add_paragraph(
        document,
        "The pooled NHANES 2007-2012 files included 30442 records. After restricting the analysis to "
        "adults, excluding pregnant participants, requiring valid UACR and common thyroid indicators, "
        "excluding thyroid-medication users, and requiring complete cohort-construction covariates, "
        "6487 participants remained (Supplementary Table S1). Survey-weighted baseline characteristics "
        "are shown in Table 1.",
    )

    document.add_heading("Association of UACR with TT4", level=2)
    add_paragraph(
        document,
        "Higher UACR was associated with higher TT4. In Model 3, each one-unit increase in natural-log "
        "UACR was associated with a 0.077-unit increase in TT4 (95% CI 0.035-0.119; P=0.001; "
        "FDR-adjusted P=0.010). Compared with UACR <30 mg/g, the adjusted TT4 difference was 0.228 "
        "(95% CI 0.032-0.425; P=0.031) for UACR 30-300 mg/g and 0.412 (95% CI 0.127-0.696; P=0.009) "
        "for UACR >=300 mg/g, with evidence of a positive clinical-category trend (P for trend=0.005) "
        "(Table 2).",
    )

    document.add_heading("TT4 robustness analyses", level=2)
    add_paragraph(
        document,
        "The UACR-TT4 association remained directionally stable across prespecified analyses. All "
        "seven natural-log UACR sensitivity models yielded positive estimates and nominal P values "
        "below 0.05. Restricted cubic spline analysis supported an overall association (P=0.00463) "
        "without evidence of non-linearity (P=0.518) (Figure 1; Supplementary Table S3). Secondary "
        "thyroid outcomes and exploratory thyroid-autoimmunity outcomes are reported in Supplementary "
        "Table S2.",
    )

    document.add_heading("Mortality outcomes", level=2)
    add_paragraph(
        document,
        "Among the 6487 discovery-cohort participants, 6484 were eligible for mortality follow-up. "
        f"The survey-weighted median follow-up was {followup_summary}. During follow-up, 897 "
        "all-cause deaths and 249 cardiovascular deaths were identified "
        "(Supplementary Table S6).",
    )
    add_paragraph(
        document,
        "Natural-log UACR was associated with both mortality outcomes. In the fully adjusted "
        "survey-weighted Cox model, the hazard ratio (HR) was 1.387 for all-cause mortality (95% CI "
        "1.285-1.498; P<0.001) and 1.331 for cardiovascular mortality (95% CI 1.161-1.526; P<0.001). "
        "TT4 showed a modest association with all-cause mortality (HR=1.065 per unit increase, 95% CI "
        "1.009-1.125; P=0.022) but not cardiovascular mortality (HR=1.108, 95% CI 0.978-1.255; "
        "P=0.108) (Table 3; Figure 2).",
    )
    add_paragraph(
        document,
        "The primary UACR mortality associations remained positive after excluding deaths within the "
        "first two years of follow-up and after restriction to euthyroid participants. TT4 estimates "
        "were less uniformly stable across sensitivity analyses and are therefore interpreted as "
        "secondary prognostic findings. Secondary joint-category analyses did not show a monotonic "
        "risk pattern. No evidence of multiplicative effect modification was observed "
        "(Supplementary Tables S7-S9).",
    )

    document.add_heading("Supplementary replication assessment and exploratory genetic analysis", level=2)
    add_paragraph(
        document,
        "The NHANES III cohort included 11302 participants after prespecified exclusions. Using "
        "harmonized covariate adjustment, the association between natural-log UACR and TT4 was not "
        "statistically replicated in NHANES III. The H3 estimate was close to zero (beta=0.003, 95% CI "
        "-0.074 to 0.080; P=0.932; n=11200), and the clinical-category analysis did not support a "
        "positive trend (P for trend=0.541) (Supplementary Tables S4 and S5).",
    )
    add_paragraph(
        document,
        "The exploratory genetic analysis could not directly evaluate UACR to TT4 because a direct "
        "TT4 GWAS was unavailable in the searchable OpenGWAS index. UACR and albuminuria each had one "
        "LD-clumped genome-wide significant instrument. The eGFR to TSH analysis was not statistically "
        "significant after false-discovery-rate correction. Full details and limitations are reported "
        "in Supplementary Tables S10-S12.",
    )

    document.add_heading("Discussion", level=1)
    add_paragraph(
        document,
        "In this nationally sampled NHANES 2007-2012 cohort, higher UACR was associated with higher "
        "TT4 after adjustment for demographic, socioeconomic, lifestyle, metabolic, kidney-function, "
        "and iodine-related covariates. The association was directionally stable across prespecified "
        "sensitivity analyses and showed no evidence of non-linearity. In the mortality extension, "
        "UACR was associated with both all-cause and cardiovascular mortality, whereas TT4 showed a "
        "modest association with all-cause mortality only.",
    )
    add_paragraph(
        document,
        "The UACR-TT4 finding may reflect shared aspects of kidney function, albumin handling, "
        "metabolic health, or thyroid hormone homeostasis. These data do not establish a mechanism. "
        "The finding should also be interpreted in light of the NHANES III analysis, in which the "
        "association was not statistically replicated. Differences in survey era, laboratory methods, "
        "population structure, covariate measurement, or residual confounding may contribute to the "
        "between-cohort heterogeneity, but the present study cannot distinguish among these "
        "possibilities.",
    )
    add_paragraph(
        document,
        "The mortality findings reinforce the established prognostic importance of albuminuria. The "
        "association of TT4 with all-cause mortality was smaller and was not observed for "
        "cardiovascular mortality. TT4 should therefore be interpreted as a secondary prognostic "
        "marker rather than as a stand-alone risk-stratification measure. Secondary joint-category "
        "analyses were descriptive: they did not show a monotonic pattern, and the effect-modification "
        "tests were not statistically significant.",
    )
    add_paragraph(
        document,
        "This study has several strengths, including a nationally sampled cohort, complex-survey "
        "methods, cycle-specific pooled weighting, multiple exposure definitions, restricted cubic "
        "spline modeling, prespecified sensitivity analyses, and prospective public-use mortality "
        "linkage. The separate NHANES III analysis is also informative because it defines an important "
        "boundary: the TT4 association was not statistically replicated in an earlier survey cohort.",
    )
    add_paragraph(
        document,
        "Several limitations should be considered. The primary UACR-TT4 analysis was cross-sectional "
        "and cannot establish temporal ordering. Complete-case criteria required common thyroid "
        "indicators and cohort-construction covariates, which may have introduced selection bias. "
        "Physical-activity completeness was required during cohort construction although physical "
        "activity was not included in the final adjustment set. Mortality analyses remain "
        "observational and may be affected by residual confounding. Cardiovascular mortality analyses "
        "included fewer events than all-cause analyses, and public-use cause-of-death categories are "
        "grouped rather than adjudicated. Finally, the exploratory genetic analysis could not directly "
        "evaluate TT4 and was limited by instrument availability.",
    )

    document.add_heading("Conclusions", level=1)
    add_paragraph(
        document,
        "Among adults in NHANES 2007-2012, higher UACR was associated with higher TT4 and with higher "
        "all-cause and cardiovascular mortality. TT4 showed a modest association with all-cause "
        "mortality but not cardiovascular mortality. Because the UACR-TT4 association was not "
        "statistically replicated in NHANES III and the genetic analysis was exploratory, additional "
        "studies are needed to determine the reproducibility and interpretation of the TT4 finding.",
    )

    document.add_heading("Acknowledgments", level=1)
    add_paragraph(document, "[INSERT ACKNOWLEDGMENTS OR STATE THAT THERE ARE NONE.]")
    document.add_heading("Author contributions", level=1)
    add_paragraph(document, AUTHOR_CONTRIBUTIONS)
    add_paragraph(document, CREDIT_STATEMENT)
    document.add_heading("Funding", level=1)
    add_paragraph(document, FUNDING_STATEMENT)
    document.add_heading("Competing interests", level=1)
    add_paragraph(document, COMPETING_INTERESTS_STATEMENT)
    document.add_heading("Data availability", level=1)
    add_paragraph(
        document,
        "NHANES, NHANES III, and public-use linked mortality data are publicly available from NCHS. "
        "OpenGWAS summary statistics used in the Supplementary exploratory genetic analysis are "
        "publicly available from OpenGWAS. Analysis scripts and derived outputs will be deposited at "
        "[INSERT PUBLIC REPOSITORY DOI OR PERSISTENT URL BEFORE SUBMISSION]. See the accompanying "
        "PLOS ONE data-availability statement for source URLs.",
    )

    document.add_heading("References", level=1)
    add_paragraph(
        document,
        "[INSERT AND FORMAT DOMAIN REFERENCES BEFORE SUBMISSION. Include UACR risk literature, "
        "kidney-thyroid literature, the 2021 CKD-EPI equation, NHANES documentation, linked mortality "
        "documentation, and STROBE reporting guidance as appropriate.]",
    )

    document.add_heading("Tables", level=1)
    document.add_heading("Table 1. Survey-weighted baseline characteristics of the NHANES 2007-2012 discovery cohort.", level=2)
    add_table(document, list(table1.columns), table1.astype(str).values.tolist(), font_size=7)
    add_paragraph(
        document,
        "Values are unweighted n (survey-weighted percentage), mean (standard error), or median "
        "(interquartile range). No between-group hypothesis tests were generated. Abbreviations: "
        "eGFR, estimated glomerular filtration rate; PIR, poverty-income ratio; TSH, thyroid-stimulating "
        "hormone; TT4, total thyroxine; UACR, urinary albumin-to-creatinine ratio.",
    )

    document.add_heading("Table 2. Association of UACR with TT4 in NHANES 2007-2012.", level=2)
    add_table(
        document,
        ["Exposure or contrast", "Model 1", "Model 2", "Model 3"],
        build_main_table_2(table2),
        font_size=8,
    )
    add_paragraph(
        document,
        "Values are beta (95% CI) unless otherwise indicated. Model 1 was unadjusted. Model 2 adjusted "
        "for age, sex, and race/ethnicity. Model 3 additionally adjusted for education, PIR, BMI, "
        "smoking, alcohol use, diabetes, hypertension, eGFR, and UIC. UACR <30 mg/g was the reference "
        "clinical category.",
    )

    document.add_heading("Table 3. Associations of natural-log UACR and TT4 with mortality.", level=2)
    add_table(
        document,
        ["Outcome", "Exposure", "Model 1 HR (95% CI)", "P", "Model 3 HR (95% CI)", "P", "Events"],
        build_main_table_3(mortality),
        font_size=8,
    )
    add_paragraph(
        document,
        "Model 1 adjusted for age, sex, and race/ethnicity. Model 3 additionally adjusted for education, "
        "PIR, BMI, smoking, alcohol use, diabetes, hypertension, eGFR, and UIC. TT4 estimates are per "
        "1-unit increase. HR, hazard ratio.",
    )

    document.add_heading("Figure legends", level=1)
    add_paragraph(
        document,
        "Figure 1. Restricted cubic spline association between UACR and TT4 in NHANES 2007-2012. "
        "The model used natural-log UACR and adjusted for age, sex, race/ethnicity, education, PIR, "
        "BMI, smoking, alcohol use, diabetes, hypertension, eGFR, and UIC. Reference lines mark UACR "
        "values of 30 and 300 mg/g. P for overall association=0.00463; P for non-linearity=0.518.",
    )
    add_paragraph(
        document,
        "Figure 2. Survey-weighted Cox model estimates for all-cause and cardiovascular mortality. "
        "Fully adjusted hazard ratios are shown for natural-log UACR and TT4 sensitivity models. "
        "TT4 is interpreted as a secondary prognostic marker.",
    )

    document.add_heading("Supporting information captions", level=1)
    for label, title in supplementary_rows():
        add_paragraph(document, f"{label}. {title}")

    output = root / "manuscript" / "PLOS_ONE_main_manuscript_draft.docx"
    document.save(output)


def make_cover_letter(root: Path) -> None:
    document = Document()
    configure_document(document, line_numbers=False)
    document.core_properties.title = f"Cover letter: {PROJECT_TITLE}"
    document.core_properties.subject = "PLOS ONE cover letter"
    document.core_properties.author = AUTHOR_NAMES
    normal = document.styles["Normal"]
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1
    normal.paragraph_format.space_after = Pt(5)
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    today = datetime.now()
    add_paragraph(document, f"{today:%B} {today.day}, {today:%Y}")
    add_paragraph(document, "Editors")
    add_paragraph(document, "PLOS ONE")
    add_paragraph(document, "Dear Editors,")
    add_paragraph(
        document,
        f'We are pleased to submit our Research Article, "{PROJECT_TITLE}," for consideration in '
        "PLOS ONE.",
    )
    add_paragraph(
        document,
        "Using survey-weighted analyses of NHANES 2007-2012, we found that higher urinary "
        "albumin-to-creatinine ratio was associated with higher total thyroxine. The association "
        "remained directionally stable across prespecified sensitivity analyses. In linked mortality "
        "analyses, natural-log UACR was associated with all-cause and cardiovascular mortality, while "
        "total thyroxine showed a modest association with all-cause mortality only.",
    )
    add_paragraph(
        document,
        "We have intentionally kept the interpretation bounded. The UACR-TT4 association was not "
        "statistically replicated in NHANES III, which is reported in the Supplementary Materials. "
        "The Supplementary exploratory genetic analysis is also presented as hypothesis-generating "
        "only. The manuscript does not make a mechanistic or causal claim.",
    )
    add_paragraph(
        document,
        "We believe the manuscript is suitable for PLOS ONE because it addresses a clinically relevant "
        "population-health question using transparent complex-survey methods, prespecified sensitivity "
        "analyses, public-use mortality linkage, and a reproducible Python and R workflow.",
    )
    add_paragraph(
        document,
        "This manuscript is original, has not been published elsewhere, and is not under consideration "
        "by another journal. All authors have approved the manuscript and agree with its submission "
        "to PLOS ONE.",
    )
    add_paragraph(document, COMPETING_INTERESTS_STATEMENT)
    add_paragraph(document, FUNDING_STATEMENT)
    add_paragraph(
        document,
        "Thank you for considering our manuscript.",
    )
    add_paragraph(document, "Sincerely,")
    add_corresponding_author_block(document)

    output = root / "manuscript" / "PLOS_ONE_cover_letter.docx"
    document.save(output)


def strobe_rows() -> list[tuple[str, str, str, str]]:
    return [
        ("1a", "Indicate the study design in the title or abstract.", "Title; Abstract", "Included in abstract; title names NHANES 2007-2012."),
        ("1b", "Provide an informative and balanced abstract summary.", "Abstract", "Included."),
        ("2", "Explain the scientific background and rationale.", "Introduction", "Included; insert final literature references."),
        ("3", "State specific objectives and prespecified hypotheses.", "Introduction", "Included."),
        ("4", "Present key elements of the study design early.", "Materials and methods: Study design", "Included."),
        ("5", "Describe the setting, locations, and relevant dates.", "Materials and methods: Study design", "Included."),
        ("6a", "Give eligibility criteria and participant selection methods.", "Materials and methods: Discovery cohort", "Included; flow in Table S1."),
        ("6b", "For matched studies, give matching criteria.", "Not applicable", "No matched design."),
        ("7", "Define outcomes, exposures, predictors, confounders, and effect modifiers.", "Materials and methods", "Included."),
        ("8", "Give data sources and measurement methods.", "Materials and methods", "Included; add final NHANES documentation references."),
        ("9", "Describe efforts to address potential sources of bias.", "Materials and methods; Discussion", "Complete-case limitations and sensitivity analyses included."),
        ("10", "Explain how the study size was arrived at.", "Results: Study population; Table S1", "Included."),
        ("11", "Explain handling of quantitative variables.", "Materials and methods", "Included."),
        ("12a", "Describe statistical methods, including confounding control.", "Materials and methods", "Included."),
        ("12b", "Describe subgroup and effect-modification methods.", "Materials and methods", "Included as secondary analyses."),
        ("12c", "Explain how missing data were addressed.", "Materials and methods: Discovery cohort", "Complete-case approach stated."),
        ("12d", "Explain loss-to-follow-up handling.", "Materials and methods: Mortality follow-up", "Public-use linkage eligibility and follow-up criteria stated."),
        ("12e", "Describe sensitivity analyses.", "Materials and methods", "Included."),
        ("13a", "Report numbers at each study stage.", "Results; Table S1", "Included."),
        ("13b", "Give reasons for non-participation at each stage.", "Table S1", "Included."),
        ("13c", "Consider use of a flow diagram.", "Table S1", "Flow table supplied; diagram may be added during final figure assembly."),
        ("14a", "Give participant characteristics and exposures.", "Table 1", "Included."),
        ("14b", "Indicate missing data for variables of interest.", "Table S1; Supplementary source table", "Complete-case flow and missingness file available."),
        ("14c", "Summarize follow-up time.", "Results: Mortality outcomes", "Survey-weighted median and interquartile range included."),
        ("15", "Report outcome events or summary measures over time.", "Results: Mortality outcomes", "897 all-cause and 249 cardiovascular deaths reported."),
        ("16a", "Give unadjusted and adjusted estimates with precision.", "Tables 2 and 3", "Included."),
        ("16b", "Report category boundaries.", "Materials and methods; table notes", "Included."),
        ("16c", "Consider translating relative risks into absolute risks.", "Not included", "Not required for the current observational scope."),
        ("17", "Report other analyses.", "Supplementary Tables S2-S12", "Included."),
        ("18", "Summarize key results with reference to objectives.", "Discussion", "Included."),
        ("19", "Discuss limitations.", "Discussion", "Included."),
        ("20", "Give a cautious overall interpretation.", "Discussion; Conclusions", "Included."),
        ("21", "Discuss generalizability.", "Discussion", "Included with survey-era boundary."),
        ("22", "Give funding source and funder role.", "Funding statement", "Included."),
    ]


def make_strobe_checklist(root: Path) -> None:
    document = Document()
    configure_document(document, line_numbers=False)
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("STROBE checklist for cohort studies")
    add_paragraph(document, f"Manuscript: {PROJECT_TITLE}")
    add_paragraph(document, f"Official checklist source: {STROBE_URL}")
    add_paragraph(document, "Page numbers should be updated after the final manuscript layout is frozen.")
    rows = [[item, recommendation, location, note] for item, recommendation, location, note in strobe_rows()]
    add_table(document, ["Item", "Recommendation", "Manuscript location", "Status or author action"], rows, font_size=8)
    output = root / "manuscript" / "PLOS_ONE_STROBE_checklist.docx"
    document.save(output)


def supplementary_rows() -> list[tuple[str, str]]:
    return [
        ("Table S1", "NHANES 2007-2012 discovery cohort exclusion flow."),
        ("Table S2", "Full discovery thyroid outcome models. TGAb and TPOAb are exploratory outcomes."),
        ("Table S3", "TT4 robustness analyses, including restricted cubic spline statistics."),
        ("Table S4", "NHANES III participant flow and diagnostic distributions."),
        ("Table S5", "Harmonized NHANES 2007-2012 and NHANES III assessment of the UACR-TT4 association."),
        ("Table S6", "Mortality linkage flow."),
        ("Table S7", "Secondary descriptive joint UACR and TT4 mortality categories."),
        ("Table S8", "Secondary mortality effect-modification tests."),
        ("Table S9", "Mortality sensitivity analyses and proportional-hazards diagnostics."),
        ("Table S10", "OpenGWAS trait selection for the exploratory genetic analysis."),
        ("Table S11", "Exploratory bidirectional genetic-analysis main results."),
        ("Table S12", "Exploratory bidirectional genetic-analysis sensitivity results."),
    ]


def make_table_lists(root: Path) -> None:
    fields = ["order", "proposed_label", "placement", "title", "source_file", "status", "submission_action"]
    main_rows = [
        {
            "order": "1",
            "proposed_label": "Table 1",
            "placement": "Main manuscript",
            "title": "Survey-weighted baseline characteristics of the NHANES 2007-2012 discovery cohort",
            "source_file": "outputs/tables/Table1_discovery_baseline_characteristics.csv",
            "status": "generated",
            "submission_action": "Retain descriptive statistics only; no between-group hypothesis tests.",
        },
        {
            "order": "2",
            "proposed_label": "Table 2",
            "placement": "Main manuscript",
            "title": "Association of UACR with TT4 in NHANES 2007-2012",
            "source_file": "outputs/tables/Table2_discovery_main_results.csv",
            "status": "embedded_in_draft",
            "submission_action": "Retain TT4 rows for natural-log UACR and clinical categories.",
        },
        {
            "order": "3",
            "proposed_label": "Table 3",
            "placement": "Main manuscript",
            "title": "Associations of natural-log UACR and TT4 with all-cause and cardiovascular mortality",
            "source_file": "outputs/tables/Table_mortality_main.csv",
            "status": "embedded_in_draft",
            "submission_action": "Retain Model 1 and fully adjusted rows for natural-log UACR and TT4.",
        },
    ]
    write_csv(root / "outputs" / "tables" / "final_main_table_list.csv", main_rows, fields)

    supplementary_sources = [
        "outputs/tables/discovery_exclusion_flow.csv",
        "outputs/tables/TableS_full_thyroid_results.csv",
        "outputs/tables/TableS_TT4_robustness.csv",
        "outputs/tables/validation_exclusion_flow.csv; outputs/tables/Table_validation_diagnostic_distribution.csv",
        "outputs/tables/Table_harmonized_discovery_validation_TT4.csv",
        "outputs/tables/mortality_linkage_flow.csv",
        "outputs/tables/Table_mortality_joint.csv",
        "outputs/tables/Table_mortality_interaction.csv",
        "outputs/tables/Table_mortality_sensitivity.csv",
        "outputs/tables/Table_MR_GWAS_selection.csv",
        "outputs/tables/Table_MR_main.csv",
        "outputs/tables/Table_MR_sensitivity.csv",
    ]
    supplementary_actions = [
        "Use as participant-flow source data.",
        "Keep TGAb and TPOAb exploratory and outside the main conclusion.",
        "Report restricted cubic spline statistics and prespecified TT4 sensitivity analyses.",
        "State that the NHANES III result was not statistically replicated.",
        "Retain harmonized H1-H3 comparisons as Supplementary evidence.",
        "Document mortality linkage and eligible follow-up counts.",
        "Present as secondary descriptive analysis; do not claim a monotonic joint-risk gradient.",
        "Report non-significant tests; do not claim effect modification.",
        "Include early-death exclusion, euthyroid restriction, per-SD TT4, and diagnostic PH checks.",
        "State unavailable FT4 and TT4 GWAS and instrument limitations.",
        "Label results as exploratory genetic analysis only.",
        "Retain as Supplementary technical detail only.",
    ]
    supplementary_rows_csv: list[dict[str, str]] = []
    for index, ((label, title), source, action) in enumerate(
        zip(supplementary_rows(), supplementary_sources, supplementary_actions), start=1
    ):
        supplementary_rows_csv.append(
            {
                "order": str(index),
                "proposed_label": label,
                "placement": "Supplementary Materials",
                "title": title.rstrip("."),
                "source_file": source,
                "status": "available",
                "submission_action": action,
            }
        )
    write_csv(
        root / "outputs" / "tables" / "final_supplementary_table_list.csv",
        supplementary_rows_csv,
        fields,
    )


def make_audit(root: Path, followup_summary: str) -> None:
    audit = f"""# Final result audit for PLOS ONE submission

## Frozen manuscript title

{PROJECT_TITLE}

## Frozen main-text storyline

1. In NHANES 2007-2012, higher natural-log UACR was associated with higher TT4.
2. The UACR-TT4 association remained positive across prespecified sensitivity analyses and showed no evidence of non-linearity.
3. Higher natural-log UACR was associated with higher all-cause and cardiovascular mortality.
4. TT4 showed a modest association with all-cause mortality but not cardiovascular mortality.
5. The NHANES III result was not statistically replicated and is placed in the Supplementary Materials.
6. The genetic module is an exploratory genetic analysis in the Supplementary Materials and does not establish causality.

## Frozen numerical results

### NHANES 2007-2012 TT4 analysis

- Discovery cohort: n=6487.
- Natural-log UACR -> TT4, fully adjusted model: beta=0.077, 95% CI 0.035-0.119, P=0.001, FDR-adjusted P=0.010.
- UACR 30-300 vs <30 mg/g -> TT4: beta=0.228, 95% CI 0.032-0.425, P=0.031.
- UACR >=300 vs <30 mg/g -> TT4: beta=0.412, 95% CI 0.127-0.696, P=0.009.
- Clinical-category trend: P=0.005.
- Restricted cubic spline: P overall=0.00463; P non-linearity=0.518.

### Mortality analysis

- Eligible follow-up: n=6484.
- Survey-weighted median follow-up: {followup_summary}.
- All-cause deaths: 897.
- Cardiovascular deaths: 249.
- Natural-log UACR -> all-cause mortality: HR=1.387, 95% CI 1.285-1.498, P<0.001.
- Natural-log UACR -> cardiovascular mortality: HR=1.331, 95% CI 1.161-1.526, P<0.001.
- TT4 -> all-cause mortality: HR=1.065, 95% CI 1.009-1.125, P=0.022.
- TT4 -> cardiovascular mortality: HR=1.108, 95% CI 0.978-1.255, P=0.108.

### Supplementary NHANES III assessment

- NHANES III cohort: n=11302; harmonized H3 n=11200.
- Natural-log UACR -> TT4, H3: beta=0.003, 95% CI -0.074 to 0.080, P=0.932.
- Clinical-category trend: P=0.541.
- Required wording: **not statistically replicated in NHANES III**.

### Supplementary exploratory genetic analysis

- Keep the genetic analysis outside the main evidence chain.
- FT4 and TT4 GWAS were unavailable in the searchable OpenGWAS index.
- UACR and albuminuria each had one LD-clumped genome-wide significant instrument.
- The TSH -> eGFR result is a single-SNP exploratory result based on a TSH protein proxy.
- Do not use the genetic results to claim a direct UACR-TT4 pathway.

## Submission artifacts generated

- `manuscript/PLOS_ONE_title_page.docx`
- `manuscript/PLOS_ONE_main_manuscript_draft.docx`
- `manuscript/PLOS_ONE_cover_letter.docx`
- `manuscript/PLOS_ONE_STROBE_checklist.docx`
- `manuscript/PLOS_ONE_data_availability_statement.md`
- `manuscript/PLOS_ONE_ethics_statement.md`
- `manuscript/PLOS_ONE_competing_interests_statement.md`
- `manuscript/PLOS_ONE_funding_statement.md`
- `outputs/tables/final_main_table_list.csv`
- `outputs/tables/final_supplementary_table_list.csv`
- `manuscript/PLOS_ONE_Supplementary_Tables.docx`
- `manuscript/PLOS_ONE_Supplementary_Tables_full.xlsx`
- `manuscript/PLOS_ONE_Supplementary_Figure_Legends.md`
- `outputs/tables/final_figure_list.csv`
- `outputs/figures/submission/Figure1_RCS_TT4.pdf`
- `outputs/figures/submission/Figure2_mortality_forest.pdf`
- `outputs/figures/submission/FigureS1_joint_mortality.pdf`
- `outputs/figures/submission/FigureS2_exploratory_MR_forest.pdf`
- `outputs/reports/PLOS_ONE_author_information_audit.md`

## Author actions required before upload

1. Confirm the study-specific CRediT contribution wording and insert acknowledgments or state that there are none.
2. Insert and format the final scientific references.
3. Deposit scripts and derived outputs in a public repository and insert the DOI or persistent URL. Exclude credentials and API tokens.
4. Assemble the Supplementary Tables file from the frozen source tables and update STROBE page numbers after final pagination.
5. Perform a final Word-layout review of tables and figure placement before upload.

## Funding record

{FUNDING_STATEMENT}

## Competing interests record

{COMPETING_INTERESTS_STATEMENT}

## Ethics record

{ETHICS_STATEMENT}

## Claim guardrails

- Do not describe NHANES III as a successful replication.
- Do not state that UACR causes higher TT4.
- Do not claim a mortality effect-modification finding for UACR and TT4.
- Do not elevate TGAb or TPOAb to primary outcomes.
- Do not place the exploratory genetic analysis in the main evidence chain.

## Official guidance checked

- PLOS ONE submission guidelines: {PLOS_GUIDELINES_URL}
- PLOS data availability policy: {PLOS_DATA_URL}
- STROBE cohort checklist: {STROBE_URL}
- NHANES ethics review information: {NHANES_ETHICS_URL}
"""
    write_text(root / "outputs" / "reports" / "final_result_audit.md", audit)


def make_author_information_audit(root: Path) -> None:
    detail_lines = []
    for name, department, institution, address, email, orcid in AUTHOR_DETAILS:
        detail_lines.extend(
            [
                f"### {name}",
                "",
                f"- Department: {department}",
                f"- Institution: {institution}",
                f"- Address: {address}",
                f"- Email: {email}",
                f"- ORCID: {orcid}",
                "",
            ]
        )
    audit = "\n".join(
        [
            "# PLOS ONE author information audit",
            "",
            "The author identities, order, affiliations, corresponding-author details, and ORCIDs were "
            "carried forward from the prior NHANES and GBD cardiovascular-risk manuscript title page.",
            "",
            "## Author line",
            "",
            AUTHOR_LINE,
            "",
            EQUAL_CONTRIBUTION,
            "",
            "## Corresponding author",
            "",
            CORRESPONDING_AUTHOR,
            CORRESPONDING_DEPARTMENT,
            CORRESPONDING_INSTITUTION,
            CORRESPONDING_ADDRESS,
            f"Email: {CORRESPONDING_EMAIL}",
            "",
            "## Full author information",
            "",
            *detail_lines,
            "## Funding",
            "",
            FUNDING_STATEMENT,
            "",
            "## Competing interests",
            "",
            COMPETING_INTERESTS_STATEMENT,
            "",
            "## Ethics",
            "",
            ETHICS_STATEMENT,
            "",
            "## Study-specific confirmation still required",
            "",
            "- Confirm the adapted CRediT contribution wording for this thyroid-UACR study.",
            "- Confirm acknowledgments.",
        ]
    )
    write_text(root / "outputs" / "reports" / "PLOS_ONE_author_information_audit.md", audit)


def validate_text_outputs(root: Path) -> None:
    targets = [
        root / "outputs" / "reports" / "final_result_audit.md",
        root / "outputs" / "reports" / "PLOS_ONE_author_information_audit.md",
        root / "outputs" / "tables" / "final_main_table_list.csv",
        root / "outputs" / "tables" / "final_supplementary_table_list.csv",
        *sorted((root / "manuscript").glob("PLOS_ONE_*.md")),
    ]
    prohibited = [
        r"externally " + r"validated",
        r"external " + r"validation",
        r"causal " + r"confirmation",
    ]
    for path in targets:
        content = path.read_text(encoding="utf-8-sig")
        for pattern in prohibited:
            if re.search(pattern, content, flags=re.IGNORECASE):
                raise RuntimeError(f"Prohibited wording '{pattern}' found in {path}.")


def main() -> None:
    root = find_project_root()
    tables = root / "outputs" / "tables"
    manuscript = root / "manuscript"
    manuscript.mkdir(parents=True, exist_ok=True)

    table1 = pd.read_csv(tables / "Table1_discovery_baseline_characteristics.csv")
    table2 = pd.read_csv(tables / "Table2_discovery_main_results.csv")
    mortality = pd.read_csv(tables / "Table_mortality_main.csv")
    descriptive = pd.read_csv(tables / "PLOS_submission_descriptive_summary.csv")
    descriptive_values = dict(zip(descriptive["metric"], descriptive["value"]))
    followup_summary = (
        f"{descriptive_values['followup_years_weighted_median']:.2f} years "
        f"(IQR {descriptive_values['followup_years_weighted_q1']:.2f}-"
        f"{descriptive_values['followup_years_weighted_q3']:.2f})"
    )

    for filename, text in statement_texts().items():
        write_text(manuscript / filename, text)
    make_table_lists(root)
    make_audit(root, followup_summary)
    make_author_information_audit(root)
    make_title_page(root)
    make_main_manuscript(root, table1, table2, mortality, followup_summary)
    make_cover_letter(root)
    make_strobe_checklist(root)
    validate_text_outputs(root)

    outputs = [
        manuscript / "PLOS_ONE_title_page.docx",
        manuscript / "PLOS_ONE_main_manuscript_draft.docx",
        manuscript / "PLOS_ONE_cover_letter.docx",
        manuscript / "PLOS_ONE_STROBE_checklist.docx",
        manuscript / "PLOS_ONE_data_availability_statement.md",
        manuscript / "PLOS_ONE_ethics_statement.md",
        manuscript / "PLOS_ONE_competing_interests_statement.md",
        manuscript / "PLOS_ONE_funding_statement.md",
        root / "outputs" / "reports" / "final_result_audit.md",
        root / "outputs" / "reports" / "PLOS_ONE_author_information_audit.md",
        tables / "final_main_table_list.csv",
        tables / "final_supplementary_table_list.csv",
    ]
    log_lines = [
        f"{datetime.now():%Y-%m-%d %H:%M:%S} | INFO | Generated PLOS ONE submission artifacts.",
        f"{datetime.now():%Y-%m-%d %H:%M:%S} | INFO | Main manuscript title: {PROJECT_TITLE.replace(chr(0x2013), '-')}",
        f"{datetime.now():%Y-%m-%d %H:%M:%S} | INFO | Embedded main tables: 3.",
        f"{datetime.now():%Y-%m-%d %H:%M:%S} | INFO | Supplementary table list rows: {len(supplementary_rows())}.",
        f"{datetime.now():%Y-%m-%d %H:%M:%S} | INFO | Text prohibited-wording scan passed.",
        f"{datetime.now():%Y-%m-%d %H:%M:%S} | INFO | NHANES III is described as not statistically replicated.",
        f"{datetime.now():%Y-%m-%d %H:%M:%S} | INFO | Genetic analysis is described as exploratory and Supplementary.",
        f"{datetime.now():%Y-%m-%d %H:%M:%S} | INFO | Author identities and affiliations carried forward from the prior NHANES and GBD manuscript.",
    ]
    for path in outputs:
        log_lines.append(
            f"{datetime.now():%Y-%m-%d %H:%M:%S} | INFO | Output: {path.relative_to(root)}"
        )
    write_text(root / "outputs" / "logs" / "11_prepare_plos_one_submission.log", "\n".join(log_lines))

    print("Generated PLOS ONE submission artifacts:")
    for path in outputs:
        print(f"- {path.relative_to(root)}")


if __name__ == "__main__":
    main()
