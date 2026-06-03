from __future__ import annotations

import csv
import re
from datetime import datetime
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


TITLE = (
    "Urinary albumin-to-creatinine ratio, total thyroxine, and mortality risk "
    "among U.S. adults: evidence from NHANES 2007-2012"
)
HEADER_FILL = "D9EAF7"
ALT_FILL = "F5F8FA"


def find_project_root(start: Path | None = None) -> Path:
    current = (start or Path.cwd()).resolve()
    for candidate in [current, *current.parents]:
        if (candidate / "config" / "analysis_plan.yaml").exists():
            return candidate
    raise RuntimeError("Could not find project root containing config/analysis_plan.yaml.")


def read_csv(root: Path, relative_path: str) -> pd.DataFrame:
    return pd.read_csv(root / relative_path)


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.rstrip() + "\n", encoding="utf-8")


def clean_text(value) -> str:
    if pd.isna(value):
        return ""
    text = str(value)
    replacements = {
        "UACR_CLINICAL_CATEGORY30-300": "UACR 30-300 vs <30 mg/g",
        "UACR_CLINICAL_CATEGORY>=300": "UACR >=300 vs <30 mg/g",
        "UACR_QUARTILEQ2": "UACR quartile Q2 vs Q1",
        "UACR_QUARTILEQ3": "UACR quartile Q3 vs Q1",
        "UACR_QUARTILEQ4": "UACR quartile Q4 vs Q1",
        "UACR_lt30__TT4_non_high": "UACR <30 + TT4 non-high",
        "UACR_ge30__TT4_non_high": "UACR >=30 + TT4 non-high",
        "UACR_lt30__TT4_high": "UACR <30 + TT4 high",
        "UACR_ge30__TT4_high": "UACR >=30 + TT4 high",
        "all_cause_mortality": "All-cause mortality",
        "cardiovascular_mortality": "Cardiovascular mortality",
        "kidney_to_thyroid": "Kidney-related trait to thyroid proxy",
        "thyroid_to_kidney": "Thyroid proxy to kidney-related trait",
    }
    return replacements.get(text, text.replace("_", " "))


def fmt_number(value, digits: int = 3) -> str:
    if pd.isna(value):
        return ""
    return f"{float(value):.{digits}f}"


def fmt_p(value) -> str:
    if pd.isna(value):
        return ""
    value = float(value)
    if value < 0.001:
        return f"{value:.2e}"
    return f"{value:.3f}"


def fmt_ci(estimate, low, high, digits: int = 3) -> str:
    if pd.isna(estimate):
        return ""
    return f"{float(estimate):.{digits}f} ({float(low):.{digits}f}-{float(high):.{digits}f})"


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(end)


def configure_document(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(8)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1
    for name, size in (("Title", 14), ("Heading 1", 11), ("Heading 2", 9)):
        style = document.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(2)
    section = document.sections[0]
    set_landscape(section)


def set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.7)
    section.page_height = Inches(8.3)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    add_page_number(section.footer.paragraphs[0])


def set_cell_margins(cell, value: int = 35) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def set_cell_text(cell, value: str, *, bold: bool = False, size: float = 7.2) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(value)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def add_table(document: Document, headers: list[str], rows: list[list[str]], *, font_size: float = 7.0) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    repeat_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, header, bold=True, size=font_size)
        shade_cell(cell, HEADER_FILL)
    for row_number, values in enumerate(rows, start=1):
        row = table.add_row()
        for index, value in enumerate(values):
            cell = row.cells[index]
            set_cell_text(cell, str(value), size=font_size)
            if row_number % 2 == 0:
                shade_cell(cell, ALT_FILL)
    document.add_paragraph()


def add_caption(document: Document, label: str, title: str, note: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(f"{label}. {title}")
    run.bold = True
    run.font.size = Pt(9)
    if note:
        paragraph = document.add_paragraph(note)
        paragraph.paragraph_format.space_after = Pt(3)
        for run in paragraph.runs:
            run.font.size = Pt(7.5)


def new_table_section(document: Document) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    set_landscape(section)


def dataframe_rows(data: pd.DataFrame, columns: list[str]) -> list[list[str]]:
    return [[clean_text(row[column]) for column in columns] for _, row in data.iterrows()]


def table_s1(root: Path) -> tuple[list[str], list[list[str]], str]:
    data = read_csv(root, "outputs/tables/discovery_exclusion_flow.csv")
    columns = ["step", "n_before", "n_excluded", "n_after", "note"]
    headers = ["Selection step", "n before", "n excluded", "n after", "Note"]
    return headers, dataframe_rows(data, columns), "Participant selection for the NHANES 2007-2012 discovery cohort."


def table_s2(root: Path) -> tuple[list[str], list[list[str]], str]:
    data = read_csv(root, "outputs/tables/TableS_full_thyroid_results.csv")
    rows = []
    for _, row in data.iterrows():
        rows.append(
            [
                clean_text(row["outcome"]),
                clean_text(row["outcome_role"]),
                clean_text(row["exposure_type"]),
                clean_text(row["contrast"]),
                clean_text(row["model"]),
                fmt_ci(row["beta"], row["ci_low"], row["ci_high"]),
                fmt_p(row["p_value"]),
                fmt_p(row["p_fdr"]),
                fmt_p(row["p_trend"]),
                str(int(row["n_model_unweighted"])),
            ]
        )
    headers = ["Outcome", "Role", "Exposure", "Contrast", "Model", "Beta (95% CI)", "P", "FDR", "P trend", "n"]
    note = (
        "Survey-weighted discovery models. TT4 is the primary thyroid outcome. TGAb and TPOAb are "
        "exploratory thyroid-autoimmunity outcomes."
    )
    return headers, rows, note


def table_s3(root: Path) -> tuple[list[str], list[list[str]], str]:
    data = read_csv(root, "outputs/tables/TableS_TT4_robustness.csv")
    rows = []
    for _, row in data.iterrows():
        rows.append(
            [
                clean_text(row["analysis_label"]),
                clean_text(row["exposure_definition"]),
                clean_text(row["contrast"]),
                fmt_ci(row["beta"], row["ci_low"], row["ci_high"]),
                fmt_p(row["p_value"]),
                fmt_p(row["p_fdr"]),
                "" if pd.isna(row["n_model_unweighted"]) else str(int(row["n_model_unweighted"])),
                clean_text(row["direction"]),
                fmt_p(row["rcs_p_overall"]),
                fmt_p(row["rcs_p_nonlinearity"]),
            ]
        )
    headers = ["Analysis", "Exposure", "Contrast", "Beta (95% CI)", "P", "FDR", "n", "Direction", "RCS P overall", "RCS P non-linearity"]
    note = "Prespecified TT4 robustness analyses. The restricted cubic spline curve is shown in Figure 1."
    return headers, rows, note


def table_s4(root: Path) -> list[tuple[str, list[str], list[list[str]], str]]:
    flow = read_csv(root, "outputs/tables/validation_exclusion_flow.csv")
    distributions = read_csv(root, "outputs/tables/Table_validation_diagnostic_distribution.csv")
    flow_headers = ["Selection step", "n before", "n excluded", "n after", "Note"]
    flow_rows = dataframe_rows(flow, ["step", "n_before", "n_excluded", "n_after", "note"])
    distribution_rows = []
    for _, row in distributions.iterrows():
        distribution_rows.append(
            [
                clean_text(row["cohort"]),
                clean_text(row["variable"]),
                clean_text(row["unit"]),
                clean_text(row["category"]),
                "" if pd.isna(row["n_nonmissing"]) else str(int(row["n_nonmissing"])),
                fmt_number(row["weighted_mean"]),
                fmt_number(row["weighted_se"]),
                fmt_number(row["weighted_median"]),
                fmt_number(row["weighted_q1"]),
                fmt_number(row["weighted_q3"]),
                fmt_number(row["weighted_percent"], 1),
            ]
        )
    distribution_headers = ["Cohort", "Variable", "Unit", "Category", "n non-missing", "Weighted mean", "SE", "Weighted median", "Q1", "Q3", "Weighted %"]
    return [
        ("Table S4A", flow_headers, flow_rows, "NHANES III participant selection flow."),
        ("Table S4B", distribution_headers, distribution_rows, "Diagnostic distribution comparison for the discovery and NHANES III cohorts."),
    ]


def table_s5(root: Path) -> tuple[list[str], list[list[str]], str]:
    data = read_csv(root, "outputs/tables/Table_harmonized_discovery_validation_TT4.csv")
    rows = []
    for _, row in data.iterrows():
        rows.append(
            [
                clean_text(row["cohort"]),
                clean_text(row["model"]),
                clean_text(row["exposure"]),
                clean_text(row["contrast"]),
                fmt_ci(row["beta"], row["ci_low"], row["ci_high"]),
                fmt_p(row["p_value"]),
                fmt_p(row["p_fdr"]),
                fmt_p(row["p_trend"]),
                str(int(row["n_model"])),
                fmt_number(row["weighted_tt4_mean"]),
                fmt_number(row["weighted_uacr_median"]),
                clean_text(row["direction"]),
            ]
        )
    headers = ["Cohort", "Model", "Exposure", "Contrast", "Beta (95% CI)", "P", "FDR", "P trend", "n", "Weighted TT4 mean", "Weighted UACR median", "Direction"]
    note = "Harmonized models omit UIC for comparability. The association was not statistically replicated in NHANES III."
    return headers, rows, note


def table_s6(root: Path) -> tuple[list[str], list[list[str]], str]:
    data = read_csv(root, "outputs/tables/mortality_linkage_flow.csv")
    headers = ["Mortality-linkage step", "n", "Note"]
    return headers, dataframe_rows(data, ["step", "n", "note"]), "Public-use linked mortality flow for the NHANES 2007-2012 discovery cohort."


def table_s7(root: Path) -> tuple[list[str], list[list[str]], str]:
    data = read_csv(root, "outputs/tables/Table_mortality_joint.csv")
    rows = []
    for _, row in data.iterrows():
        rows.append(
            [
                clean_text(row["outcome"]),
                clean_text(row["model"]),
                clean_text(row["high_definition"]),
                clean_text(row["group"]),
                fmt_ci(row["hr"], row["ci_low"], row["ci_high"]),
                fmt_p(row["p_value"]),
                str(int(row["n_model"])),
                str(int(row["events"])),
                str(int(row["group_n"])),
                str(int(row["group_events"])),
            ]
        )
    headers = ["Outcome", "Model", "TT4-high definition", "Group", "HR (95% CI)", "P", "Model n", "Events", "Group n", "Group events"]
    note = "Secondary descriptive analysis only. The joint groups did not support a monotonic combined-risk claim."
    return headers, rows, note


def table_s8(root: Path) -> tuple[list[str], list[list[str]], str]:
    data = read_csv(root, "outputs/tables/Table_mortality_interaction.csv")
    rows = []
    for _, row in data.iterrows():
        rows.append(
            [
                clean_text(row["outcome"]),
                clean_text(row["model"]),
                clean_text(row["interaction_type"]),
                clean_text(row["high_definition"]),
                fmt_ci(row["interaction_hr"], row["ci_low"], row["ci_high"]),
                fmt_p(row["p_for_interaction"]),
                str(int(row["n_model"])),
                str(int(row["events"])),
            ]
        )
    headers = ["Outcome", "Model", "Effect-modification term", "Definition", "Interaction HR (95% CI)", "P for interaction", "n", "Events"]
    note = "Secondary tests. No statistically significant multiplicative effect modification was observed in fully adjusted models."
    return headers, rows, note


def table_s9(root: Path) -> list[tuple[str, list[str], list[list[str]], str]]:
    data = read_csv(root, "outputs/tables/Table_mortality_sensitivity.csv")
    cox = data[data["analysis_type"] == "survey_weighted_cox"].copy()
    diagnostics = data[data["analysis_type"] == "ph_diagnostic_unweighted"].copy()
    cox_rows = []
    for _, row in cox.iterrows():
        cox_rows.append(
            [
                clean_text(row["scenario"]),
                clean_text(row["outcome"]),
                clean_text(row["exposure"]),
                clean_text(row["contrast"]),
                fmt_ci(row["hr"], row["ci_low"], row["ci_high"]),
                fmt_p(row["p_value"]),
                str(int(row["n_model"])),
                str(int(row["events"])),
                clean_text(row["direction"]),
            ]
        )
    diagnostic_rows = []
    for _, row in diagnostics.iterrows():
        diagnostic_rows.append(
            [
                clean_text(row["outcome"]),
                clean_text(row["exposure"]),
                clean_text(row["ph_term"]),
                fmt_p(row["ph_p_value"]),
                clean_text(row["note"]),
            ]
        )
    return [
        (
            "Table S9A",
            ["Scenario", "Outcome", "Exposure", "Contrast", "HR (95% CI)", "P", "n", "Events", "Direction"],
            cox_rows,
            "Survey-weighted Cox sensitivity analyses.",
        ),
        (
            "Table S9B",
            ["Outcome", "Exposure", "PH diagnostic term", "Diagnostic P", "Note"],
            diagnostic_rows,
            "Diagnostic non-weighted proportional-hazards checks using Model 3 covariates.",
        ),
    ]


def table_s10(root: Path) -> tuple[list[str], list[list[str]], str]:
    data = read_csv(root, "outputs/tables/Table_MR_GWAS_selection.csv")
    columns = ["trait_key", "status", "selected_id", "selected_trait", "population", "sample_size", "year", "unit", "n_instruments", "note"]
    headers = ["Trait", "Status", "OpenGWAS ID", "Selected trait", "Population", "Sample size", "Year", "Unit", "Instruments", "Note"]
    return headers, dataframe_rows(data, columns), "OpenGWAS trait selection for the Supplementary exploratory genetic analysis."


def table_s11(root: Path) -> tuple[list[str], list[list[str]], str]:
    data = read_csv(root, "outputs/tables/Table_MR_main.csv")
    rows = []
    for _, row in data.iterrows():
        rows.append(
            [
                clean_text(row["direction"]),
                clean_text(row["exposure_trait"]),
                clean_text(row["outcome_trait"]),
                clean_text(row["method"]),
                "" if pd.isna(row["nsnp"]) else str(int(row["nsnp"])),
                fmt_ci(row["beta"], row["ci_low"], row["ci_high"], 4),
                fmt_p(row["p_value"]),
                fmt_p(row["p_fdr"]),
                clean_text(row["status"]),
                clean_text(row["note"]),
            ]
        )
    headers = ["Direction", "Exposure", "Outcome", "Method", "SNPs", "Beta (95% CI)", "P", "FDR", "Status", "Note"]
    note = "Exploratory genetic analysis only. Direct TT4 analyses were unavailable in the searchable OpenGWAS index."
    return headers, rows, note


def table_s12(root: Path) -> list[tuple[str, list[str], list[list[str]], str]]:
    data = read_csv(root, "outputs/tables/Table_MR_sensitivity.csv")
    completed = data[(data["status"] == "complete") & (data["method"] != "leave-one-out IVW")].copy()
    status = (
        data.groupby(["method", "status"], dropna=False)
        .size()
        .reset_index(name="rows")
        .sort_values(["method", "status"])
    )
    completed_rows = []
    for _, row in completed.iterrows():
        completed_rows.append(
            [
                clean_text(row["analysis_id"]),
                clean_text(row["exposure_trait"]),
                clean_text(row["outcome_trait"]),
                clean_text(row["method"]),
                clean_text(row["detail"]),
                "" if pd.isna(row["nsnp"]) else str(int(row["nsnp"])),
                fmt_number(row["beta"], 4),
                fmt_p(row["p_value"]),
                clean_text(row["note"]),
            ]
        )
    status_rows = [[clean_text(row["method"]), clean_text(row["status"]), str(int(row["rows"]))] for _, row in status.iterrows()]
    return [
        (
            "Table S12A",
            ["Analysis", "Exposure", "Outcome", "Method", "Detail", "SNPs", "Beta", "P", "Note"],
            completed_rows,
            "Completed exploratory genetic sensitivity analyses excluding individual leave-one-out rows.",
        ),
        (
            "Table S12B",
            ["Method", "Status", "Rows in complete workbook"],
            status_rows,
            "Method-status inventory. The complete 444-row technical output, including leave-one-out rows, is retained in the accompanying XLSX workbook.",
        ),
    ]


def supplementary_specs(root: Path):
    return [
        ("Table S1", "NHANES 2007-2012 discovery cohort exclusion flow", [table_s1(root)]),
        ("Table S2", "Full discovery thyroid outcome models", [table_s2(root)]),
        ("Table S3", "TT4 robustness analyses", [table_s3(root)]),
        ("Table S4", "NHANES III participant flow and diagnostic distributions", table_s4(root)),
        ("Table S5", "Harmonized NHANES 2007-2012 and NHANES III assessment of the UACR-TT4 association", [table_s5(root)]),
        ("Table S6", "Mortality linkage flow", [table_s6(root)]),
        ("Table S7", "Secondary descriptive joint UACR and TT4 mortality categories", [table_s7(root)]),
        ("Table S8", "Secondary mortality effect-modification tests", [table_s8(root)]),
        ("Table S9", "Mortality sensitivity analyses and proportional-hazards diagnostics", table_s9(root)),
        ("Table S10", "OpenGWAS trait selection for the exploratory genetic analysis", [table_s10(root)]),
        ("Table S11", "Exploratory bidirectional genetic-analysis main results", [table_s11(root)]),
        ("Table S12", "Exploratory bidirectional genetic-analysis sensitivity results", table_s12(root)),
    ]


def make_docx(root: Path) -> Path:
    document = Document()
    configure_document(document)
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Supplementary Tables")
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(TITLE).bold = True
    document.add_paragraph(
        "This document contains reader-facing Supplementary Tables. Complete machine-readable "
        "source tables, including all 444 exploratory genetic sensitivity rows, are retained in "
        "the accompanying XLSX workbook."
    )
    document.add_paragraph(
        "Interpretation boundary: NHANES III was not statistically replicated. The genetic module "
        "is exploratory and does not establish causality. Joint mortality categories are secondary "
        "descriptive analyses."
    )
    document.add_heading("Contents", level=1)
    index_rows = [[label, title] for label, title, _ in supplementary_specs(root)]
    add_table(document, ["Table", "Title"], index_rows, font_size=7.5)

    for label, title, sub_tables in supplementary_specs(root):
        new_table_section(document)
        document.add_heading(f"{label}. {title}", level=1)
        for sub_index, spec in enumerate(sub_tables):
            if len(spec) == 4:
                section_label, headers, rows, note = spec
            else:
                headers, rows, note = spec
                section_label = f"{label}{chr(ord('A') + sub_index)}"
            if len(sub_tables) > 1:
                document.add_heading(section_label, level=2)
            document.add_paragraph(note)
            font_size = 6.4 if len(headers) >= 10 else 6.8 if len(headers) >= 8 else 7.2
            add_table(document, headers, rows, font_size=font_size)

    document.add_heading("Supplementary Figure Legends", level=1)
    figure_legends = [
        (
            "Figure S1",
            "Secondary descriptive joint UACR and TT4 mortality categories. Hazard ratios and 95% "
            "confidence intervals are from fully adjusted survey-weighted Cox models. TT4 high was "
            "defined as the weighted highest quartile. The analysis is descriptive and did not "
            "support a monotonic combined-risk claim."
        ),
        (
            "Figure S2",
            "Supplementary exploratory genetic-analysis IVW estimates. Square symbols denote the "
            "multi-SNP eGFR-to-TSH estimate; circles denote single-SNP estimates. Direct TT4 genetic "
            "analyses were unavailable in the searchable OpenGWAS index. This analysis does not "
            "establish causality."
        ),
    ]
    for label, legend in figure_legends:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{label}. ").bold = True
        paragraph.add_run(legend)

    output = root / "manuscript" / "PLOS_ONE_Supplementary_Tables.docx"
    document.save(output)
    return output


def auto_width(worksheet) -> None:
    for column_cells in worksheet.columns:
        length = max(len(str(cell.value or "")) for cell in column_cells)
        width = min(max(length + 2, 10), 55)
        worksheet.column_dimensions[get_column_letter(column_cells[0].column)].width = width


def add_sheet(workbook: Workbook, name: str, data: pd.DataFrame) -> None:
    worksheet = workbook.create_sheet(title=name[:31])
    worksheet.append(list(data.columns))
    for row in data.itertuples(index=False, name=None):
        worksheet.append(["" if pd.isna(value) else value for value in row])
    for cell in worksheet[1]:
        cell.font = Font(name="Arial", bold=True, color="000000")
        cell.fill = PatternFill("solid", fgColor=HEADER_FILL)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for row in worksheet.iter_rows(min_row=2):
        for cell in row:
            cell.font = Font(name="Arial", size=9)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    worksheet.freeze_panes = "A2"
    worksheet.auto_filter.ref = worksheet.dimensions
    auto_width(worksheet)


def make_xlsx(root: Path) -> Path:
    workbook = Workbook()
    info = workbook.active
    info.title = "README"
    readme_rows = [
        ["Supplementary Tables workbook", TITLE],
        ["Purpose", "Complete machine-readable tables supporting the reader-facing Supplementary Tables DOCX."],
        ["Interpretation boundary", "NHANES III was not statistically replicated. Genetic analyses are exploratory only."],
        ["Generated", f"{datetime.now():%Y-%m-%d %H:%M:%S}"],
    ]
    for row in readme_rows:
        info.append(row)
    for cell in info[1]:
        cell.font = Font(name="Arial", bold=True)
    info.column_dimensions["A"].width = 26
    info.column_dimensions["B"].width = 120
    for row in info.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.font = Font(name="Arial", size=10, bold=cell.row == 1)

    sheets = [
        ("S1_discovery_flow", "outputs/tables/discovery_exclusion_flow.csv"),
        ("S2_full_thyroid", "outputs/tables/TableS_full_thyroid_results.csv"),
        ("S3_TT4_robustness", "outputs/tables/TableS_TT4_robustness.csv"),
        ("S4a_validation_flow", "outputs/tables/validation_exclusion_flow.csv"),
        ("S4b_validation_diag", "outputs/tables/Table_validation_diagnostic_distribution.csv"),
        ("S5_harmonized_TT4", "outputs/tables/Table_harmonized_discovery_validation_TT4.csv"),
        ("S6_mortality_flow", "outputs/tables/mortality_linkage_flow.csv"),
        ("S7_mortality_joint", "outputs/tables/Table_mortality_joint.csv"),
        ("S8_mortality_interaction", "outputs/tables/Table_mortality_interaction.csv"),
        ("S9_mortality_sensitivity", "outputs/tables/Table_mortality_sensitivity.csv"),
        ("S10_MR_selection", "outputs/tables/Table_MR_GWAS_selection.csv"),
        ("S11_MR_main", "outputs/tables/Table_MR_main.csv"),
        ("S12_MR_sensitivity_full", "outputs/tables/Table_MR_sensitivity.csv"),
    ]
    for sheet_name, relative_path in sheets:
        add_sheet(workbook, sheet_name, read_csv(root, relative_path))
    output = root / "manuscript" / "PLOS_ONE_Supplementary_Tables_full.xlsx"
    workbook.save(output)
    return output


def make_figure_legends(root: Path) -> Path:
    text = """# PLOS ONE Supplementary Figure Legends

**Figure S1. Secondary descriptive joint UACR and TT4 mortality categories.** Hazard ratios and 95% confidence intervals are from fully adjusted survey-weighted Cox models. TT4 high was defined as the weighted highest quartile. The analysis is descriptive and did not support a monotonic combined-risk claim.

**Figure S2. Supplementary exploratory genetic-analysis IVW estimates.** Square symbols denote the multi-SNP eGFR-to-TSH estimate; circles denote single-SNP estimates. Direct TT4 genetic analyses were unavailable in the searchable OpenGWAS index. This analysis does not establish causality.
"""
    output = root / "manuscript" / "PLOS_ONE_Supplementary_Figure_Legends.md"
    write_text(output, text)
    return output


def update_table_list(root: Path) -> None:
    path = root / "outputs" / "tables" / "final_supplementary_table_list.csv"
    data = pd.read_csv(path)
    data["status"] = "formatted_docx_and_full_xlsx"
    data["reader_facing_docx"] = "manuscript/PLOS_ONE_Supplementary_Tables.docx"
    data["complete_workbook"] = "manuscript/PLOS_ONE_Supplementary_Tables_full.xlsx"
    data.to_csv(path, index=False, encoding="utf-8-sig")


def make_audit(root: Path, docx_path: Path, xlsx_path: Path) -> Path:
    text = f"""# PLOS ONE Supplementary Tables formatting audit

## Outputs

- Reader-facing DOCX: `{docx_path.relative_to(root).as_posix()}`
- Complete machine-readable workbook: `{xlsx_path.relative_to(root).as_posix()}`
- Supplementary figure legends: `manuscript/PLOS_ONE_Supplementary_Figure_Legends.md`

## Formatting decisions

- Tables are arranged in manuscript order from S1 to S12.
- Reader-facing columns use shortened labels and combined estimate-with-confidence-interval fields.
- Repeated table headers and compact Arial typography are applied throughout the DOCX.
- Alternating row shading improves scanning without changing any values.
- Complete CSV-derived rows are preserved in the XLSX workbook.
- Table S12 is concise in the DOCX; its complete 444-row technical output, including leave-one-out results, is preserved in the XLSX workbook.

## Interpretation boundaries retained

- NHANES III was not statistically replicated.
- TGAb and TPOAb remain exploratory thyroid-autoimmunity outcomes.
- Joint mortality categories are secondary descriptive analyses.
- Genetic analyses are exploratory and do not establish causality.

## Author review

- Open the DOCX in Word and inspect page breaks before upload.
- Confirm that the XLSX workbook is uploaded as a Supporting Information file if the journal submission form permits it.
"""
    output = root / "outputs" / "reports" / "PLOS_ONE_supplementary_tables_audit.md"
    write_text(output, text)
    return output


def validate(root: Path, outputs: Iterable[Path]) -> None:
    for path in outputs:
        if not path.exists() or path.stat().st_size == 0:
            raise RuntimeError(f"Missing or empty output: {path}")
    patterns = [
        r"externally " + r"validated",
        r"external " + r"validation",
        r"causal " + r"confirmation",
    ]
    for path in [
        root / "outputs" / "tables" / "final_supplementary_table_list.csv",
        root / "manuscript" / "PLOS_ONE_Supplementary_Figure_Legends.md",
        root / "outputs" / "reports" / "PLOS_ONE_supplementary_tables_audit.md",
    ]:
        text = path.read_text(encoding="utf-8-sig")
        for pattern in patterns:
            if re.search(pattern, text, re.I):
                raise RuntimeError(f"Prohibited wording '{pattern}' found in {path}")


def main() -> None:
    root = find_project_root()
    docx_path = make_docx(root)
    xlsx_path = make_xlsx(root)
    legends_path = make_figure_legends(root)
    update_table_list(root)
    audit_path = make_audit(root, docx_path, xlsx_path)
    outputs = [docx_path, xlsx_path, legends_path, audit_path]
    validate(root, outputs)

    log = root / "outputs" / "logs" / "13_prepare_plos_supplementary_tables.log"
    write_text(
        log,
        "\n".join(
            [
                f"{datetime.now():%Y-%m-%d %H:%M:%S} | INFO | Generated reader-facing Supplementary Tables DOCX.",
                f"{datetime.now():%Y-%m-%d %H:%M:%S} | INFO | Generated complete Supplementary Tables XLSX workbook.",
                f"{datetime.now():%Y-%m-%d %H:%M:%S} | INFO | Preserved full Table S12 technical output in workbook.",
                f"{datetime.now():%Y-%m-%d %H:%M:%S} | INFO | Updated final supplementary table list.",
            ]
        ),
    )
    print("Generated:")
    for path in outputs:
        print(f"- {path.relative_to(root)}")


if __name__ == "__main__":
    main()
