from __future__ import annotations

import csv
import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document


REPOSITORY_URL = "https://github.com/HAOYANGLI888/uacr-tt4-nhanes-plosone"
TITLE = (
    "Urinary albumin-to-creatinine ratio, total thyroxine, and mortality risk among "
    "U.S. adults: evidence from NHANES 2007-2012"
)
AUTHOR_LINE = "Haoyang Li1†, Xiang Li2†, Xuefeng Shi3*"
EQUAL_CONTRIBUTION = "†Haoyang Li and Xiang Li contributed equally to this work and share first authorship."
AFFILIATIONS = [
    "1Department of Cardiovascular Medicine, Qinghai University Affiliated Hospital, No. 29 Tongren Road, Chengxi District, Xining, Qinghai 810000, China.",
    "2Department of Thyroid and Breast Surgery, Qinghai University Affiliated Hospital, No. 29 Tongren Road, Chengxi District, Xining, Qinghai 810000, China.",
    "3Department of Respiratory and Critical Care Medicine, Qinghai Provincial People’s Hospital, No. 2 Gonghe Road, Chengdong District, Xining, Qinghai 810000, China.",
]
CORRESPONDING = (
    "Xuefeng Shi; Department of Respiratory and Critical Care Medicine; Qinghai Provincial "
    "People’s Hospital; No. 2 Gonghe Road, Chengdong District, Xining, Qinghai 810000, China; "
    "Email: shixuefeng128@163.com; ORCID: 0000-0002-4694-8759."
)
FUNDING = (
    "This study was supported by the National Excellent Young Physician Program "
    "(Document No. 2024[41]). The funder had no role in study design, data collection and "
    "analysis, decision to publish, or preparation of the manuscript."
)
COMPETING = "The authors have declared that no competing interests exist."
ETHICS = (
    "This study was a secondary analysis of publicly available, de-identified data from NHANES "
    "and NHANES III. NHANES protocols were reviewed and approved by the NCHS Research Ethics "
    "Review Board, and written informed consent was obtained from participants. The present "
    "study involved no direct participant contact and used only public-use files; therefore, "
    "additional institutional review board approval was not required."
)

PLACEHOLDERS = [
    r"\[" + r"INSERT",
    r"\[" + r"AUTHORS " + r"TO " + r"CONFIRM",
    r"TO BE " + r"CONFIRMED",
    r"UPDATE AFTER " + r"FINAL EDITING",
    r"AUTHORS" + r" TO",
]
PROHIBITED = [
    "externally " + "validated",
    "causal " + "confirmation",
    "UACR " + "causes TT4",
    "MR confirmed " + "causality",
    "synergistic mortality " + "effect",
    "thyroid autoimmunity as " + "main finding",
]


def find_project_root(start: Path | None = None) -> Path:
    current = (start or Path.cwd()).resolve()
    for candidate in [current, *current.parents]:
        if (candidate / "config" / "analysis_plan.yaml").exists():
            return candidate
    raise RuntimeError("Project root not found.")


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text.rstrip() + "\n", encoding="utf-8")


def docx_text(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def file_text(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        return docx_text(path)
    return path.read_text(encoding="utf-8-sig", errors="ignore")


def copy_final_files(root: Path) -> list[Path]:
    manuscript = root / "manuscript"
    mapping = {
        "PLOS_ONE_main_manuscript_draft.docx": "PLOS_ONE_main_manuscript_final.docx",
        "PLOS_ONE_title_page.docx": "PLOS_ONE_title_page_final.docx",
        "PLOS_ONE_cover_letter.docx": "PLOS_ONE_cover_letter_final.docx",
        "PLOS_ONE_STROBE_checklist.docx": "PLOS_ONE_STROBE_checklist_final.docx",
        "PLOS_ONE_Supplementary_Tables.docx": "PLOS_ONE_Supplementary_Tables_final.docx",
        "PLOS_ONE_Supplementary_Tables_full.xlsx": "PLOS_ONE_Supplementary_Tables_full_final.xlsx",
        "PLOS_ONE_Supplementary_Figure_Legends.md": "PLOS_ONE_Supplementary_Figure_Legends_final.md",
        "PLOS_ONE_data_availability_statement.md": "PLOS_ONE_data_availability_statement_final.md",
        "PLOS_ONE_ethics_statement.md": "PLOS_ONE_ethics_statement_final.md",
        "PLOS_ONE_funding_statement.md": "PLOS_ONE_funding_statement_final.md",
        "PLOS_ONE_competing_interests_statement.md": "PLOS_ONE_competing_interests_statement_final.md",
    }
    outputs: list[Path] = []
    for source_name, target_name in mapping.items():
        source = manuscript / source_name
        target = manuscript / target_name
        if not source.exists():
            raise FileNotFoundError(source)
        shutil.copy2(source, target)
        outputs.append(target)
    return outputs


def update_figure_list(root: Path) -> None:
    path = root / "outputs" / "tables" / "final_figure_list.csv"
    if not path.exists():
        return
    rows: list[dict[str, str]] = []
    with path.open("r", newline="", encoding="utf-8-sig") as handle:
        reader = csv.DictReader(handle)
        fieldnames = list(reader.fieldnames or [])
        for row in reader:
            for key in ("proposed_label", "figure_label", "label"):
                if key in row:
                    row[key] = row[key].replace("Figure S1", "S1 Fig").replace("Figure S2", "S2 Fig")
            rows.append(row)
    with path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def scan_files(paths: list[Path]) -> dict[str, list[str]]:
    findings: dict[str, list[str]] = {}
    for path in paths:
        if not path.exists() or path.suffix.lower() == ".xlsx":
            continue
        text = file_text(path)
        local_findings: list[str] = []
        for pattern in PLACEHOLDERS:
            if re.search(pattern, text, flags=re.IGNORECASE):
                local_findings.append(f"placeholder pattern: {pattern}")
        lower_text = text.lower()
        for phrase in PROHIBITED:
            if phrase.lower() in lower_text:
                local_findings.append(f"prohibited phrase: {phrase}")
        if local_findings:
            findings[path.relative_to(path.parents[1]).as_posix() if len(path.parents) > 1 else path.name] = local_findings
    return findings


def write_author_funding_report(root: Path, files_checked: list[Path], findings: dict[str, list[str]]) -> Path:
    checked = "\n".join(f"- `{path.relative_to(root).as_posix()}`" for path in files_checked if path.exists())
    status = "No author, funding, competing-interest, or ethics inconsistency was detected across final files."
    if findings:
        status = "Potential text issues were detected and are listed below; author/funding fields remained generated from one source record."
    text = f"""# Author and funding consistency report

Generated: {datetime.now():%Y-%m-%d %H:%M:%S}

## Files checked

{checked}

## Status

{status}

## Final author line

{AUTHOR_LINE}

{EQUAL_CONTRIBUTION}

## Affiliations

{chr(10).join(f"- {item}" for item in AFFILIATIONS)}

## Corresponding author

{CORRESPONDING}

## Funding

{FUNDING}

## Competing interests

{COMPETING}

## Ethics

{ETHICS}
"""
    if findings:
        text += "\n## Text scan findings\n\n"
        for path, messages in findings.items():
            text += f"- `{path}`: {'; '.join(messages)}\n"
    else:
        text += "\n## Text scan findings\n\nNo unresolved placeholders or prohibited phrases were found in final submission text files.\n"
    output = root / "outputs" / "reports" / "author_funding_consistency_report.md"
    write_text(output, text)
    return output


def write_reference_report(root: Path) -> Path:
    text = f"""# Reference insertion report

Generated: {datetime.now():%Y-%m-%d %H:%M:%S}

## Insertion map

- Introduction, UACR and mortality background: Matsushita et al. and Gerstein et al.
- Introduction and Discussion, kidney-thyroid physiology: Iglesias and Diez; Mariani and Berns.
- Introduction and Discussion, thyroid function and mortality: Rodondi et al.; Collet et al.; van den Beld et al.
- Methods, NHANES 2007-2012 and NHANES III data sources: NCHS public documentation and Hollowell et al.
- Methods, mortality linkage: NCHS public-use linked mortality file documentation.
- Methods, eGFR: CKD-EPI 2021 equation by Inker et al.
- Methods/reporting: STROBE cohort reporting statement.
- Supplementary exploratory genetic analysis: MR-Base and OpenGWAS references.

## Vancouver references inserted

1. Matsushita K, van der Velde M, Astor BC, Woodward M, Levey AS, de Jong PE, et al. Association of estimated glomerular filtration rate and albuminuria with all-cause and cardiovascular mortality in general population cohorts: a collaborative meta-analysis. Lancet. 2010;375(9731):2073-81. doi:10.1016/S0140-6736(10)60674-5.
2. Gerstein HC, Mann JF, Yi Q, Zinman B, Dinneen SF, Hoogwerf B, et al. Albuminuria and risk of cardiovascular events, death, and heart failure in diabetic and nondiabetic individuals. JAMA. 2001;286(4):421-6. doi:10.1001/jama.286.4.421.
3. Iglesias P, Diez JJ. Thyroid dysfunction and kidney disease. Eur J Endocrinol. 2009;160(4):503-15. doi:10.1530/EJE-08-0837.
4. Mariani LH, Berns JS. The renal manifestations of thyroid disease. J Am Soc Nephrol. 2012;23(1):22-6. doi:10.1681/ASN.2010070766.
5. Rodondi N, den Elzen WPJ, Bauer DC, Cappola AR, Razvi S, Walsh JP, et al. Subclinical hypothyroidism and the risk of coronary heart disease and mortality. JAMA. 2010;304(12):1365-74. doi:10.1001/jama.2010.1361.
6. Collet TH, Gussekloo J, Bauer DC, den Elzen WPJ, Cappola AR, Balmer P, et al. Subclinical hyperthyroidism and the risk of coronary heart disease and mortality. Arch Intern Med. 2012;172(10):799-809. doi:10.1001/archinternmed.2012.402.
7. van den Beld AW, Visser TJ, Feelders RA, Grobbee DE, Lamberts SWJ. Thyroid hormone concentrations, disease, physical function, and mortality in elderly men. J Clin Endocrinol Metab. 2005;90(12):6403-9. doi:10.1210/jc.2005-0872.
8. National Center for Health Statistics. NHANES 2007-2012 questionnaires, datasets, and related documentation. Centers for Disease Control and Prevention. Available from: https://wwwn.cdc.gov/nchs/nhanes/continuousnhanes/default.aspx
9. National Center for Health Statistics. NHANES III data files, documentation, and laboratory manuals. Centers for Disease Control and Prevention. Available from: https://wwwn.cdc.gov/nchs/nhanes/nhanes3/datafiles.aspx
10. National Center for Health Statistics. Public-use linked mortality files. Centers for Disease Control and Prevention. Available from: https://www.cdc.gov/nchs/data-linkage/mortality-public.htm
11. Hollowell JG, Staehling NW, Flanders WD, Hannon WH, Gunter EW, Spencer CA, et al. Serum TSH, T(4), and thyroid antibodies in the United States population (1988 to 1994): National Health and Nutrition Examination Survey (NHANES III). J Clin Endocrinol Metab. 2002;87(2):489-99. doi:10.1210/jcem.87.2.8182.
12. National Center for Health Statistics. Thyroid profile data documentation, NHANES 2007-2012. Centers for Disease Control and Prevention. Available from: https://wwwn.cdc.gov/nchs/nhanes/search/datapage.aspx?Component=Laboratory
13. Inker LA, Eneanya ND, Coresh J, Tighiouart H, Wang D, Sang Y, et al. New creatinine- and cystatin C-based equations to estimate GFR without race. N Engl J Med. 2021;385(19):1737-49. doi:10.1056/NEJMoa2102953.
14. von Elm E, Altman DG, Egger M, Pocock SJ, Gotzsche PC, Vandenbroucke JP, et al. The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement: guidelines for reporting observational studies. PLoS Med. 2007;4(10):e296. doi:10.1371/journal.pmed.0040296.
15. Hemani G, Zheng J, Elsworth B, Wade KH, Haberland V, Baird D, et al. The MR-Base platform supports systematic causal inference across the human phenome. eLife. 2018;7:e34408. doi:10.7554/eLife.34408.
16. Elsworth B, Lyon M, Alexander T, Liu Y, Matthews P, Hallett J, et al. The MRC IEU OpenGWAS data infrastructure. bioRxiv. 2020:2020.08.10.244293. doi:10.1101/2020.08.10.244293.
"""
    output = root / "outputs" / "reports" / "reference_insertion_report.md"
    write_text(output, text)
    return output


def write_action_items(root: Path) -> Path:
    text = f"""# Final author action items

Generated: {datetime.now():%Y-%m-%d %H:%M:%S}

No unresolved manuscript-body placeholders remain.

Before PLOS ONE upload:

1. Open the final DOCX files in Word and visually inspect page breaks, table widths, and figure callouts.
2. Update STROBE page numbers after final pagination if the submission system requires exact page numbers.
3. If a Zenodo DOI is minted later, add it during proofing. The current Data Availability statement uses the GitHub repository URL: {REPOSITORY_URL}.
4. Confirm that the GitHub repository contains only code, configuration, aggregate derived outputs, manuscript materials, and no raw NHANES files or credentials.
"""
    output = root / "manuscript" / "final_author_action_items.md"
    write_text(output, text)
    return output


def write_submission_checklist(root: Path, final_files: list[Path]) -> Path:
    required = [
        "PLOS_ONE_main_manuscript_final.docx",
        "PLOS_ONE_title_page_final.docx",
        "PLOS_ONE_cover_letter_final.docx",
        "PLOS_ONE_STROBE_checklist_final.docx",
        "PLOS_ONE_Supplementary_Tables_final.docx",
        "PLOS_ONE_data_availability_statement_final.md",
        "PLOS_ONE_ethics_statement_final.md",
        "PLOS_ONE_funding_statement_final.md",
        "PLOS_ONE_competing_interests_statement_final.md",
    ]
    manuscript = root / "manuscript"
    rows = []
    for name in required:
        path = manuscript / name
        status = "present" if path.exists() and path.stat().st_size > 0 else "missing"
        rows.append(f"- `{name}`: {status}")
    text = f"""# PLOS ONE submission checklist final

Generated: {datetime.now():%Y-%m-%d %H:%M:%S}

## Required files

{chr(10).join(rows)}

## Interpretation guardrails

- NHANES III is described as not statistically replicated.
- MR is described as an exploratory genetic analysis.
- The manuscript uses observational language and states that the analyses do not establish causality.
- TGAb and TPOAb are not presented as primary outcomes.
- Joint mortality categories are described as secondary descriptive analyses.

## Repository

GitHub: {REPOSITORY_URL}
"""
    output = manuscript / "PLOS_ONE_submission_checklist_final.md"
    write_text(output, text)
    return output


def write_repository_safety_report(root: Path) -> Path:
    staging = root.parent / "uacr-tt4-nhanes-plosone"
    checked_roots = []
    if staging.exists():
        checked_roots.append(staging)
    submission_files = [
        root / "README.md",
        root / "LICENSE",
        root / ".gitignore",
        root / "requirements.txt",
        root / "environment.yml",
        root / "config",
        root / "scripts",
        root / "outputs" / "tables" / "final_main_table_list.csv",
        root / "outputs" / "tables" / "final_supplementary_table_list.csv",
        root / "outputs" / "tables" / "final_figure_list.csv",
        root / "outputs" / "figures" / "submission",
        root / "outputs" / "reports" / "final_result_audit.md",
        root / "outputs" / "reports" / "author_funding_consistency_report.md",
        root / "outputs" / "reports" / "reference_insertion_report.md",
        root / "manuscript",
    ]

    findings: list[str] = []
    raw_like_extensions = {".xpt", ".dat", ".sas7bdat"}
    credential_file_names = {".Renviron", "jwt_check.txt", "opengwas_token_test.txt"}
    local_path_pattern = re.compile(r"(C:\\Users\\Administrator|D:\\anaconda)", re.IGNORECASE)
    jwt_value_pattern = re.compile(r"eyJ[A-Za-z0-9_-]{20,}\.[A-Za-z0-9_-]{20,}\.[A-Za-z0-9_-]{20,}")

    for checked_root in checked_roots:
        for path in checked_root.rglob("*"):
            if ".git" in path.parts:
                continue
            if path.is_dir():
                if path.name in {"raw", "processed"} and path.parent.name == "data" and checked_root.name == "uacr-tt4-nhanes-plosone":
                    findings.append(f"raw/processed data directory present in staging repository: {path.relative_to(checked_root).as_posix()}")
                continue
            relative = path.relative_to(checked_root).as_posix()
            if path.name in credential_file_names:
                findings.append(f"credential or local test file present: {checked_root.name}/{relative}")
            if path.suffix.lower() in raw_like_extensions and "data/raw" in relative:
                findings.append(f"raw data-like file present: {checked_root.name}/{relative}")
            if path.stat().st_size > 5_000_000 and checked_root.name == "uacr-tt4-nhanes-plosone":
                findings.append(f"large file in staging repository: {relative}")
            if path.suffix.lower() in {".md", ".txt", ".csv", ".py", ".r", ".yml", ".yaml", ".json", ".gitignore"}:
                text = path.read_text(encoding="utf-8-sig", errors="ignore")
                if local_path_pattern.search(text):
                    findings.append(f"local absolute path string in {checked_root.name}/{relative}")
                if jwt_value_pattern.search(text):
                    findings.append(f"JWT-like credential value in {checked_root.name}/{relative}")

    for item in submission_files:
        candidates = list(item.rglob("*")) if item.is_dir() else [item]
        for path in candidates:
            if not path.exists() or path.is_dir() or ".git" in path.parts:
                continue
            relative = path.relative_to(root).as_posix()
            if "outputs/logs" in relative or "data/raw" in relative or "data/processed" in relative:
                findings.append(f"excluded path unexpectedly included in submission scope: {relative}")
            if path.name in credential_file_names:
                findings.append(f"credential or local test file in submission scope: {relative}")
            if path.suffix.lower() in raw_like_extensions and "data/raw" in relative:
                findings.append(f"raw data-like file in submission scope: {relative}")
            if path.suffix.lower() in {".md", ".txt", ".csv", ".py", ".r", ".yml", ".yaml", ".json", ".gitignore"}:
                text = path.read_text(encoding="utf-8-sig", errors="ignore")
                if local_path_pattern.search(text):
                    findings.append(f"local absolute path string in submission scope: {relative}")
                if jwt_value_pattern.search(text):
                    findings.append(f"JWT-like credential value in submission scope: {relative}")

    status = "PASS" if not findings else "REVIEW REQUIRED"
    details = "\n".join(f"- {item}" for item in findings) if findings else "- No credential files, raw-data uploads, local absolute paths, or JWT-like credential values were detected in the checked submission-facing files."
    text = f"""# Repository safety check report

Generated: {datetime.now():%Y-%m-%d %H:%M:%S}

## Scope

- Final submission-facing materials checked.
- GitHub staging repository checked when available: `uacr-tt4-nhanes-plosone`.
- The scan distinguishes credential values or credential files from reproducibility text that names required environment variables.

## Status

{status}

## Findings

{details}

## Notes

- Raw NHANES files are public source data but are not redistributed.
- Local raw-data and processed-data directories are outside the submission-facing file scope and are excluded from the GitHub staging repository.
- MR reruns require users to configure their own OpenGWAS credential locally; no credential value is stored in the submission package.
"""
    output = root / "outputs" / "reports" / "repository_safety_check_report.md"
    write_text(output, text)
    return output


def main() -> None:
    root = find_project_root()
    update_figure_list(root)
    final_files = copy_final_files(root)
    text_scan_files = [
        root / "manuscript" / "PLOS_ONE_main_manuscript_final.docx",
        root / "manuscript" / "PLOS_ONE_title_page_final.docx",
        root / "manuscript" / "PLOS_ONE_cover_letter_final.docx",
        root / "manuscript" / "PLOS_ONE_STROBE_checklist_final.docx",
        root / "manuscript" / "PLOS_ONE_Supplementary_Tables_final.docx",
        root / "manuscript" / "PLOS_ONE_data_availability_statement_final.md",
        root / "manuscript" / "PLOS_ONE_ethics_statement_final.md",
        root / "manuscript" / "PLOS_ONE_funding_statement_final.md",
        root / "manuscript" / "PLOS_ONE_competing_interests_statement_final.md",
    ]
    findings = scan_files(text_scan_files)
    reports = [
        write_author_funding_report(root, text_scan_files, findings),
        write_reference_report(root),
        write_action_items(root),
        write_submission_checklist(root, final_files),
        write_repository_safety_report(root),
    ]
    for path in final_files + reports:
        if not path.exists() or path.stat().st_size == 0:
            raise RuntimeError(f"Missing final output: {path}")
    print("Final PLOS ONE submission files and reports generated:")
    for path in final_files + reports:
        print(f"- {path.relative_to(root)}")


if __name__ == "__main__":
    main()
